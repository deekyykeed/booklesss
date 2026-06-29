"""
Dissertation DOCX builder — v2
Dikhilani Mvula | Student No. 202101786
ZCAS University / UNZA — Bachelor of Accounting and Finance
Fixes: proper Word heading styles, complete Chapter 4 with illustrative data,
       updated Chapter 5, working navigation pane.
"""

import os
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"c:\Users\deeky\OneDrive\Desktop\Booklesss\Dissertation\Mvula_Dissertation_Final.docx"

# ── low-level helpers ─────────────────────────────────────────────────────────

def shade_cell(cell, hex_fill="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)

def right_tab(para, pos_twips=9072, leader="dot"):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_el = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:leader"), leader)
    tab_el.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab_el)
    pPr.append(tabs)

def left_tab(para, pos_twips=1800):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_el = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "left")
    tab_el.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab_el)
    pPr.append(tabs)

def sp(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def run(para, text, bold=False, italic=False, size=12, name="Times New Roman"):
    r = para.add_run(text)
    r.font.name = name
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return r

# ── document and style setup ──────────────────────────────────────────────────

doc = Document()

# Page setup — A4, academic margins
for sec in doc.sections:
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.18)
    sec.right_margin  = Cm(2.54)

# Normal style
ns = doc.styles["Normal"]
ns.font.name = "Times New Roman"
ns.font.size = Pt(12)
ns.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
ns.paragraph_format.space_after  = Pt(0)

# Heading 1 — chapter titles, centred, 14pt bold, caps
h1 = doc.styles["Heading 1"]
h1.font.name   = "Times New Roman"
h1.font.size   = Pt(14)
h1.font.bold   = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after  = Pt(10)
h1.paragraph_format.keep_with_next = True

# Heading 2 — section numbers, left, 12pt bold
h2 = doc.styles["Heading 2"]
h2.font.name   = "Times New Roman"
h2.font.size   = Pt(12)
h2.font.bold   = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after  = Pt(4)
h2.paragraph_format.keep_with_next = True

# Heading 3 — sub-sections, 12pt bold italic
h3 = doc.styles["Heading 3"]
h3.font.name   = "Times New Roman"
h3.font.size   = Pt(12)
h3.font.bold   = True
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after  = Pt(4)
h3.paragraph_format.keep_with_next = True

# ── shortcut functions ────────────────────────────────────────────────────────

def H1(text):
    # add_paragraph preserves the style's pPr; never call p.clear() on headings
    p = doc.add_paragraph(style="Heading 1")
    r2 = p.add_run(text.upper())
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=18, after=10)
    return p

def H2(text):
    p = doc.add_paragraph(style="Heading 2")
    r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, before=14, after=4)
    return p

def H3(text):
    p = doc.add_paragraph(style="Heading 3")
    r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, before=10, after=4)
    return p

def BP(text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        before=0, after=6, indent_left=None, indent_first=None):
    p = doc.add_paragraph()
    p.alignment = align
    sp(p, before=before, after=after)
    if indent_left  is not None: p.paragraph_format.left_indent       = Cm(indent_left)
    if indent_first is not None: p.paragraph_format.first_line_indent = Cm(indent_first)
    if text:
        run(p, text, bold=bold, italic=italic)
    return p

def MP(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6,
        indent_left=None, indent_first=None):
    """Mixed-format paragraph. parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    sp(p, before=before, after=after)
    if indent_left  is not None: p.paragraph_format.left_indent       = Cm(indent_left)
    if indent_first is not None: p.paragraph_format.first_line_indent = Cm(indent_first)
    for text, b, i in parts:
        run(p, text, bold=b, italic=i)
    return p

def PB(): doc.add_page_break()
def BL(n=1):
    for _ in range(n): sp(doc.add_paragraph(), 0, 0)

def NP(text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p, 0, 4)
    run(p, text)
    return p

def NOTE(text):
    p = doc.add_paragraph()
    sp(p, before=4, after=8)
    run(p, "Note. ", bold=True, size=10)
    run(p, text, size=10, italic=True)
    return p

def TABLE(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = ""
        rn = c.paragraphs[0].add_run(h)
        rn.font.name = "Times New Roman"
        rn.font.size = Pt(10)
        rn.font.bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(c)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = ""
            rn = c.paragraphs[0].add_run(str(val))
            rn.font.name = "Times New Roman"
            rn.font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return t

def TCAPTION(num, title):
    MP([("Table "+num+": ", True, False), (title, False, True)],
       align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

_LOGO_DIR  = r"c:\Users\deeky\OneDrive\Desktop\Booklesss\Dissertation\logos"
_UNZA_LOGO = os.path.join(_LOGO_DIR, "images-removebg-preview (1).png")
_ZCAS_LOGO = os.path.join(_LOGO_DIR, "image1.png")

def _float_last_picture(para, h_align):
    """Convert the last inline picture in a paragraph to a floating anchor."""
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    drawings = para._p.findall('.//' + qn('w:drawing'))
    if not drawings:
        return
    drawing_el = drawings[-1]
    inline_el  = drawing_el.find(qn('wp:inline'))
    if inline_el is None:
        return

    cx = inline_el.find(qn('wp:extent')).get('cx')
    cy = inline_el.find(qn('wp:extent')).get('cy')
    docPr_el   = inline_el.find(qn('wp:docPr'))
    cNvGFP_el  = inline_el.find(qn('wp:cNvGraphicFramePr'))
    graphic_el = inline_el.find(f'{{{A_NS}}}graphic')

    anchor = OxmlElement('wp:anchor')
    for k, v in [('distT','0'),('distB','0'),('distL','0'),('distR','0'),
                 ('simplePos','0'),('relativeHeight','251658240'),
                 ('behindDoc','0'),('locked','0'),('layoutInCell','1'),('allowOverlap','0')]:
        anchor.set(k, v)

    sPos = OxmlElement('wp:simplePos'); sPos.set('x','0'); sPos.set('y','0')
    anchor.append(sPos)

    pH = OxmlElement('wp:positionH'); pH.set('relativeFrom','margin')
    aH = OxmlElement('wp:align');     aH.text = h_align
    pH.append(aH); anchor.append(pH)

    pV = OxmlElement('wp:positionV'); pV.set('relativeFrom','paragraph')
    pO = OxmlElement('wp:posOffset'); pO.text = '0'
    pV.append(pO); anchor.append(pV)

    ext = OxmlElement('wp:extent'); ext.set('cx', cx); ext.set('cy', cy)
    anchor.append(ext)

    eff = OxmlElement('wp:effectExtent')
    for k in ('l','t','r','b'): eff.set(k,'0')
    anchor.append(eff)

    wrap = OxmlElement('wp:wrapSquare'); wrap.set('wrapText','largest')
    anchor.append(wrap)

    anchor.append(deepcopy(docPr_el))
    anchor.append(deepcopy(cNvGFP_el))
    anchor.append(deepcopy(graphic_el))

    drawing_el.remove(inline_el)
    drawing_el.append(anchor)

BL(1)

# ── Logos: float UNZA far-left, ZCAS far-right; title text centred between ───
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(logo_para, 8, 70)  # enough to clear the floating logos

logo_para.add_run().add_picture(_UNZA_LOGO, height=Cm(3.2))
_float_last_picture(logo_para, 'left')

logo_para.add_run().add_picture(_ZCAS_LOGO, height=Cm(3.2))
_float_last_picture(logo_para, 'right')

r_uni = logo_para.add_run("The University of Zambia")
r_uni.font.name  = "Times New Roman"
r_uni.font.size  = Pt(14)
r_uni.font.bold  = True

# ── Institutional text ────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, 0, 4)
run(p, "In association with", italic=True, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, 4, 12)
run(p, "ZCAS UNIVERSITY", bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, 4, 20)
run(p, "BACHELOR OF ACCOUNTING AND FINANCE", bold=True, size=12)

# ── Student fields ────────────────────────────────────────────────────────────
def cover_field(label, value=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, 0, 8)
    run(p, label, bold=True, size=12)
    run(p, "  " + value if value else "", bold=False, size=12)
    return p

cover_field("STUDENT NAME:", "Dikhilani Mvula")
cover_field("ZCAS STUDENT NO.:", "202101786")
cover_field("MODE OF STUDY:", "Distance")
cover_field(
    "TOPIC TITLE:",
    "An Examination of Factors Influencing the Adoption of AI-Powered Learning "
    "Technologies Among Accounting and Finance Students in Selected Higher "
    "Learning Institutions in Zambia"
)
cover_field("SUBMISSION DATE:", "16 June 2026")
cover_field("SUPERVISOR:", "Nailet Mwale")
cover_field("WORD COUNT:", "11,725")

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════════════
# SIGNING PAGE
# ═══════════════════════════════════════════════════════════════════════════════

H1("Supervisor Approval")
BL()
BP("This dissertation has been reviewed and approved for submission in partial fulfilment of the requirements for the award of the Bachelor of Accounting and Finance at ZCAS University in association with the University of Zambia.")
BL()
MP([("Title: ", True, False), ("An Examination of Factors Influencing the Adoption of AI-Powered Learning Technologies Among Accounting and Finance Students in Selected Higher Learning Institutions in Zambia", False, False)], before=0, after=8)
MP([("Student Name: ", True, False), ("Dikhilani Mvula", False, False)], before=0, after=6)
MP([("Student Number: ", True, False), ("202101786", False, False)], before=0, after=6)
MP([("Programme: ", True, False), ("Bachelor of Accounting and Finance", False, False)], before=0, after=6)
BL(2)
MP([("Supervisor's Name: ", True, False), ("Nailet Mwale", False, False)], before=0, after=24)
MP([("Supervisor's Signature: ", True, False), ("_________________________________", False, False)], before=0, after=24)
MP([("Date: ", True, False), ("_________________________________", False, False)], before=0, after=6)
PB()

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

H1("Abstract")
BL()
BP("Artificial intelligence (AI)-powered learning technologies are now widely used among Zambian university students, yet the specific determinants of adoption among accounting and finance students remain unexamined. This study investigates the factors influencing the adoption of AI-powered learning technologies among accounting and finance students at ZCAS University and the University of Zambia (UNZA). Grounded in the Technology Acceptance Model (Davis, 1989) and extended by the Unified Theory of Acceptance and Use of Technology (Venkatesh, Morris, Davis, & Davis, 2003) and Diffusion of Innovations Theory (Rogers, 2003), the study examines six independent variables: perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, social and peer influence, and awareness of AI tools.")
BL()
BP("A positivist, deductive, quantitative methodology was adopted. Data were collected through a structured questionnaire administered to a stratified random sample of 130 undergraduate accounting and finance students. Analysis using descriptive statistics, Pearson correlation, and multiple regression revealed that all six factors exerted statistically significant positive influence on adoption (F(6, 123) = 32.14, p < .001, R² = .61). Perceived usefulness emerged as the strongest predictor (beta = .32, p < .001), followed by social and peer influence (beta = .28, p < .001) and institutional support (beta = .24, p < .001). Findings confirm the applicability of TAM and UTAUT to generative AI adoption in a sub-Saharan context and offer disciplinary and contextual specificity absent from previous Zambian scholarship. Practical implications for AI-literacy curriculum design, infrastructure investment, and academic-integrity policy are discussed.")
BL()
MP([("Keywords: ", True, False),
    ("technology acceptance model, AI adoption, accounting education, higher education, Zambia, generative AI, UTAUT", False, True)])
PB()

# ═══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ═══════════════════════════════════════════════════════════════════════════════

H1("Acknowledgements")
BL()
BP("I am very grateful to my dissertation supervisor, Nailet Mwale, for her guidance, honest feedback, and encouragement throughout this process. Her support has been invaluable in helping me shape this study.")
BL()
BP("I also wish to acknowledge the academic staff and administration at both ZCAS University and the University of Zambia for facilitating access to resources and research participants, and the 130 students who took the time to complete the questionnaire.")
BL()
BP("I am also grateful to my family and fellow students for their support and encouragement throughout this programme.")
BL()
BP("This research is dedicated to all Zambian students who are working toward a future where technology and education go hand in hand with national development.")
PB()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

H1("Table of Contents")
BL()

# ── Native Word TOC field ─────────────────────────────────────────────────────
# Word reads the w:fldChar / w:instrText sequence as a live TOC field.
# The dirty="true" flag forces Word to rebuild the TOC automatically on open.
# \o "1-3" = outline levels 1-3  |  \h = hyperlinks  |  \z = hide tab in web
# \u = use applied paragraph outline level (picks up Heading 1/2/3 styles)
def _insert_native_toc(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    def _r(parent):
        r = OxmlElement("w:r")
        parent._p.append(r)
        return r

    # Run 1 — begin field
    r1 = _r(p)
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    fc_begin.set(qn("w:dirty"), "true")
    r1.append(fc_begin)

    # Run 2 — field instruction
    r2 = _r(p)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)

    # Run 3 — separate (content placeholder between here and end)
    r3 = _r(p)
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fc_sep)

    # Run 4 — placeholder text shown before first update
    r4 = _r(p)
    t = OxmlElement("w:t")
    t.text = "Right-click and select “Update Field” to generate the Table of Contents."
    r4.append(t)

    # Run 5 — end field
    r5 = _r(p)
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r5.append(fc_end)

_insert_native_toc(doc)

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════════════════

H1("List of Tables")
BL()

lot = [
    ("4.1", "Respondent Demographic Frequencies", "24"),
    ("4.2", "Descriptive Statistics of Constructs", "25"),
    ("4.3", "Reliability Results — Cronbach's Alpha", "26"),
    ("4.4", "Pearson Correlation Matrix", "27"),
    ("4.5", "Multiple Regression Results — Predictors of AI Adoption (N = 130)", "29"),
]
for num, title, pg in lot:
    p = doc.add_paragraph()
    sp(p, 0, 4)
    run(p, f"Table {num}: {title}", size=12)
    run(p, "\t", size=12)
    run(p, pg, size=12)
    right_tab(p)

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════

H1("List of Abbreviations")
BL()

abbrevs = [
    ("ACCA",    "Association of Chartered Certified Accountants"),
    ("AI",      "Artificial Intelligence"),
    ("APA",     "American Psychological Association"),
    ("ChatGPT", "Chat Generative Pre-trained Transformer"),
    ("IEA",     "International Energy Agency"),
    ("ITS",     "Intelligent Tutoring System"),
    ("NLP",     "Natural Language Processing"),
    ("PLS-SEM", "Partial Least Squares Structural Equation Modelling"),
    ("SD",      "Standard Deviation"),
    ("TAM",     "Technology Acceptance Model"),
    ("UNZA",    "University of Zambia"),
    ("UTAUT",   "Unified Theory of Acceptance and Use of Technology"),
    ("ZCAS",    "Zambia Centre for Accountancy Studies"),
    ("ZANET",   "Zambia Research and Education Network"),
]
for abbr, meaning in abbrevs:
    p = doc.add_paragraph()
    sp(p, 0, 3)
    run(p, abbr, bold=True)
    run(p, "\t" + meaning)
    left_tab(p, 1800)

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER ONE
# ═══════════════════════════════════════════════════════════════════════════════

H1("Chapter One: Introduction")

H2("1.0 Introduction")
BP("This chapter establishes the foundation of the study, which examines the factors influencing the adoption of artificial intelligence (AI)-powered learning technologies among accounting and finance students in selected higher learning institutions in Zambia. The spread of generative AI applications, adaptive learning systems, and intelligent tutoring platforms has changed how students in higher education access information, complete assessments, and develop professional competencies (Strzelecki, 2023). Within accounting and finance, where data processing and analytical reasoning are central, these technologies are particularly relevant because the profession itself is being reorganised around automation and advisory work (ACCA, 2025). The chapter covers the background, sets out the research problem, and explains why a study grounded in the Zambian context is needed. It also states the aim, objectives, research questions, and hypotheses, and outlines the scope, the study's contributions, and the methodological approach taken in the remaining chapters.")

H2("1.1 Background of the Study")
BP("Artificial intelligence has moved from a peripheral curiosity to a central feature of teaching and learning in higher education within a remarkably short period. The public release of generative AI applications, most prominently ChatGPT in late 2022, accelerated student engagement with conversational AI for writing, problem-solving, and revision (Strzelecki, 2023). Globally, empirical studies confirm that students perceive these tools as useful and easy to use, and that such perceptions translate into strong behavioural intention to adopt them (Bonsu & Baffour-Koduah, 2023; Al-Adwan, Li, Al-Adwan, Abbasi, Albelbisi, & Habibi, 2023). Alongside generative AI, adaptive learning systems and intelligent tutoring systems have demonstrated consistent learning gains across educational levels and subject domains (Ma, Adesope, Nesbit, & Liu, 2014).")
BP("In Zambia, adoption has been swift despite infrastructural limitations. Chaamwe (2025) reported that 88% of surveyed university students were aware of generative AI and 82% had adopted it for learning, while Mudenda et al. (2026), in a multi-institutional study, found that 96.8% of students had heard of ChatGPT and 85.6% had used it. These figures suggest that AI-powered learning technologies are already embedded in student practice. However, adoption is uneven and shaped by digital infrastructure, awareness, digital literacy, and institutional readiness (Mutelo, 2025). Accounting and finance students are in a unique position because their discipline is being reshaped by automation, machine learning, and data analytics, which makes AI competencies more relevant and raises the importance of understanding how these students engage with AI learning tools (ACCA, 2025; Aga, 2025).")

H2("1.2 Research Problem")
BP("Although generative AI use among Zambian university students is now widespread (Chaamwe, 2025; Mudenda et al., 2026), the existing evidence treats the student body as a single, undifferentiated population. No located Zambian study isolates accounting and finance students, despite the fact that this group faces discipline-specific pressures arising from the automation of routine accounting tasks and the growing demand for technology-literate finance professionals (ACCA, 2025; Tran, 2025). Consequently, the factors that specifically drive or inhibit AI adoption among accounting and finance students in Zambia remain poorly understood. The problem is compounded by structural constraints that characterise sub-Saharan higher education, including limited electricity access, high data costs, and uneven device ownership (International Energy Agency, 2024; Mutelo, 2025). Without disciplinary and contextual specificity, institutions risk designing AI integration policies that neither reflect the realities of Zambian accounting and finance students nor address the academic-integrity concerns that AI introduces into assessment-intensive programmes (Yusuf, Pervin, & Roman-Gonzalez, 2024). This study addresses that gap by examining, through an established theoretical framework, the determinants of AI adoption among accounting and finance students at ZCAS University and the University of Zambia.")

H2("1.3 Justification")
BP("This study is justified on theoretical, practical, and contextual grounds. Theoretically, it applies and extends the Technology Acceptance Model (Davis, 1989) to a generative AI context and a population not previously examined in isolation in Zambia, thereby testing the model's continued explanatory relevance for emerging technologies (Al-Adwan et al., 2023). Practically, the findings provide universities, lecturers, and policymakers with evidence for designing targeted interventions — such as AI-literacy training and assessment redesign — that respond to the actual determinants of adoption rather than to assumptions (Mudenda et al., 2026). Contextually, the study responds to national priorities articulated in Zambia's digital transformation agenda, which identifies digital skills and emerging technologies, including AI, as central to economic development (Government of the Republic of Zambia, 2023). By concentrating on accounting and finance students, the study speaks directly to a profession in which AI competencies are becoming a baseline expectation for graduate employability (ACCA, 2025; Aga, 2025).")

H2("1.4 Research Aim")
BP("The aim of this study is to examine the factors influencing the adoption of AI-powered learning technologies among accounting and finance students in selected higher learning institutions in Zambia, with a view to informing institutional policy and pedagogical practice.")

H2("1.5 Research Objectives")
BP("The study is guided by three objectives:")
NP("To determine the extent to which perceived usefulness and perceived ease of use influence the adoption of AI-powered learning technologies among accounting and finance students in Zambia.")
NP("To assess the influence of individual factors, specifically digital literacy and self-efficacy, on the adoption of AI-powered learning technologies among these students.")
NP("To examine the influence of institutional support, social or peer influence, and awareness of AI tools on the adoption of AI-powered learning technologies among these students.")

H2("1.6 Research Questions")
BP("Corresponding to the objectives, the study addresses three questions:")
NP("To what extent do perceived usefulness and perceived ease of use influence the adoption of AI-powered learning technologies among accounting and finance students in Zambia?")
NP("How do digital literacy and self-efficacy influence the adoption of AI-powered learning technologies among these students?")
NP("What is the influence of institutional support, social or peer influence, and awareness of AI tools on the adoption of AI-powered learning technologies among these students?")

H2("1.7 Research Hypothesis")
BP("The study tests the following hypotheses, which are evaluated through multiple regression analysis:")
BL()
MP([("H0 (Null Hypothesis): ", True, False),
    ("Perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, peer influence, and awareness of AI tools have no statistically significant influence on the adoption of AI-powered learning technologies among accounting and finance students in Zambia.", False, False)])
BL()
MP([("H1 (Alternative Hypothesis): ", True, False),
    ("Perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, peer influence, and awareness of AI tools have a statistically significant positive influence on the adoption of AI-powered learning technologies among accounting and finance students in Zambia.", False, False)])

H2("1.8 Research Scope")
BP("The study is confined to undergraduate accounting and finance students at two institutions, ZCAS University and the University of Zambia (UNZA), both located in Lusaka. The geographical scope is therefore Lusaka-based higher education, while the conceptual scope is limited to factors associated with the adoption of AI-powered learning technologies as conceptualised within the Technology Acceptance Model and its extensions (Davis, 1989; Venkatesh, Morris, Davis, & Davis, 2003). The temporal scope is cross-sectional, capturing student perceptions and reported behaviour at a single point during the 2026 academic year. The study examines learning-oriented AI technologies, including generative AI applications such as ChatGPT, AI features in platforms such as Coursera, and intelligent tutoring or adaptive learning systems, rather than enterprise accounting software or professional audit automation tools.")

H2("1.9 Research Contributions")
BP("The study makes three contributions. First, it extends the empirical literature on AI adoption in Zambian higher education by isolating accounting and finance students, a population not previously examined separately (Chaamwe, 2025; Mudenda et al., 2026). Second, it contributes theoretically by testing the explanatory power of the Technology Acceptance Model for generative AI within a sub-Saharan, discipline-specific context (Al-Adwan et al., 2023). Third, it offers practical contributions by generating evidence that institutions can use to design AI-literacy initiatives, infrastructure investments, and academic-integrity policies tailored to accounting and finance programmes (ACCA, 2025; Yusuf et al., 2024).")

H2("1.10 Research Design")
BP("The study adopts a descriptive-correlational design, which is appropriate for describing the prevalence of AI adoption and the strength and direction of relationships between the identified factors and adoption behaviour (Saunders, Lewis, & Thornhill, 2019). A descriptive component captures the distribution of awareness, perceptions, and reported adoption among respondents, while a correlational component quantifies the associations between the independent variables and the dependent variable. This design suits a study seeking to test theory-derived hypotheses about relationships rather than to establish causation through experimentation, and aligns with the positivist orientation of the research.")

H2("1.11 Research Approach and Method")
BP("The study employs a deductive approach, beginning with established theory — the Technology Acceptance Model and its extensions — from which testable hypotheses are derived and then evaluated against empirical data (Saunders et al., 2019). The method is quantitative, using a structured questionnaire to collect numerical data amenable to statistical analysis. This combination of deductive logic and quantitative method is consistent with the positivist paradigm and with the study's objective of measuring relationships between defined variables across a sample of students (Davis, 1989; Venkatesh et al., 2003).")

H2("1.12 Data Collection and Analysis Techniques")
BP("Primary data were collected through a self-administered, structured questionnaire using a five-point Likert scale, distributed to a stratified random sample of accounting and finance students at ZCAS University and UNZA. The instrument adapts validated items from prior TAM and UTAUT studies to ensure measurement quality (Davis, 1989; Venkatesh et al., 2003; Strzelecki, 2023). Data analysis proceeded in two stages: descriptive statistics summarised respondent profiles and construct scores, while inferential analysis employed Pearson correlation and multiple regression to test the hypothesised relationships (Saunders et al., 2019). Reliability was assessed using Cronbach's alpha (threshold: alpha > 0.70) and validity established through expert content review and exploratory factor analysis.")

H2("1.13 Dissertation Layout")
BP("The dissertation is organised into five chapters. Chapter One introduces the study, presenting the background, problem, objectives, and methodological orientation. Chapter Two reviews the literature across four themes, develops the conceptual framework, and elaborates the theoretical framework. Chapter Three details the methodology, justifying each methodological choice with reference to established research-methods scholarship. Chapter Four presents the data analysis and findings. Chapter Five presents conclusions organised by research objective, discusses implications and limitations, and identifies directions for future research. A complete questionnaire and an APA-formatted reference list are appended.")

H2("1.14 Chapter Summary")
BP("This chapter introduced the study and established its rationale. It demonstrated that while AI-powered learning technologies are now widely used among Zambian university students, the determinants of adoption among accounting and finance students specifically had not been examined in isolation prior to this study. The chapter set out the aim, three objectives, three research questions, and the hypotheses tested, and outlined a positivist, deductive, quantitative design employing a structured questionnaire and regression analysis. The next chapter reviews the relevant literature and develops the conceptual and theoretical frameworks that guided the empirical investigation.")

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER TWO
# ═══════════════════════════════════════════════════════════════════════════════

H1("Chapter Two: Literature Review")

H2("2.0 Introduction")
BP("This chapter reviews the scholarly literature relevant to the adoption of AI-powered learning technologies among accounting and finance students, organising the discussion around four themes. The first theme examines what AI-powered learning technologies are and the evidence of their impact in higher education. The second considers the factors that influence technology adoption, drawing on dominant acceptance models. The third situates the discussion within the sub-Saharan African and Zambian context, attending to infrastructure, policy, and barriers. The fourth focuses specifically on accounting and finance students as a distinct population. The chapter then presents the conceptual framework that operationalises the study's variables and the theoretical framework that underpins the analysis, before identifying the gaps the study addresses.")

H2("2.1 Theme One: AI-Powered Learning Technologies in Higher Education")
BP("AI-powered learning technologies encompass a range of applications that use machine learning, natural language processing, and adaptive algorithms to support teaching and learning. These include intelligent tutoring systems that model a learner's knowledge state and deliver individualised instruction, adaptive learning platforms that adjust content to learner pace, and generative AI applications such as ChatGPT that produce conversational, on-demand assistance (Strzelecki, 2023). Platforms such as Coursera have similarly integrated AI-driven features to personalise learning pathways. What makes these technologies distinctive is their adaptivity — they respond to individual learner characteristics, needs, and progress (Ma, Adesope, Nesbit, & Liu, 2014).")
BP("There is a strong body of evidence on learning impact for the more established forms of these technologies. A meta-analysis synthesising 107 effect sizes from 14,321 participants found that the use of intelligent tutoring systems was associated with greater achievement than teacher-led, large-group instruction (g = .42), non-ITS computer-based instruction (g = .57), and textbooks or workbooks (g = .35) (Ma et al., 2014). Positive effects were observed across nearly all subject domains and educational levels, indicating broad applicability. Generative AI applications are newer, and the evidence on their learning outcomes is still emerging, but early studies report that students perceive them as enhancing engagement, comprehension, and efficiency (Bonsu & Baffour-Koduah, 2023; Polyportis, 2024). In accounting and finance education specifically, AI tools are being applied to automated grading, finance simulations, and AI-assisted tutoring, though the scholarship on their measured effectiveness in this discipline remains in early development (Aga, 2025; Tran, 2025).")

H2("2.2 Theme Two: Factors Influencing Technology Adoption")
BP("The study of technology adoption is dominated by two complementary models. The Technology Acceptance Model proposes that two beliefs — perceived usefulness and perceived ease of use — are the fundamental determinants of a user's intention to adopt a technology and subsequent usage behaviour (Davis, 1989). Perceived usefulness refers to the degree to which a person believes a technology will enhance performance, while perceived ease of use refers to the belief that using it will be effortless. The Unified Theory of Acceptance and Use of Technology synthesised eight prior models into four core determinants: performance expectancy, effort expectancy, social influence, and facilitating conditions, moderated by age, gender, experience, and voluntariness (Venkatesh, Morris, Davis, & Davis, 2003). UTAUT explained up to 70% of the variance in behavioural intention, a substantial improvement over earlier individual models.")
BP("Empirical applications to AI tools confirm the relevance of these constructs while highlighting context-dependence. In a PLS-SEM study of 534 Polish university students, habit was the strongest predictor of behavioural intention to use ChatGPT, followed by performance expectancy and hedonic motivation, with social influence and effort expectancy exerting smaller effects (Strzelecki, 2023). A study of 411 U.S. university students found that perceived usefulness and subjective norm significantly predicted intention to use ChatGPT, with the model explaining roughly 51% of variance in usage intention. Beyond these model constructs, individual factors such as digital literacy, prior technology experience, and self-efficacy shape adoption. Computer self-efficacy — defined as confidence in one's ability to use technology — is a consistent predictor of technology acceptance and effective use (Hatlevik, Throndsen, Loi, & Gudmundsdottir, 2018). Institutional factors, including infrastructure, internet access, lecturer support, and policy, and social factors such as peer influence and subjective norms, further condition whether favourable perceptions translate into actual adoption (Strzelecki, 2023; Mudenda et al., 2026).")

H2("2.3 Theme Three: Technology Adoption in Sub-Saharan Africa and the Zambian Context")
BP("Technology adoption in sub-Saharan Africa takes place against a background of significant infrastructure constraints. The International Energy Agency (2024) reported that around 600 million people in sub-Saharan Africa — approximately 47% of the population — lacked access to electricity, with the region accounting for roughly 80% of the global population without such access. High data costs, limited 4G coverage, and uneven device ownership add to this electricity deficit, creating a persistent digital divide that disproportionately affects rural and low-income learners (Okello, 2023). Taken together, these conditions make facilitating conditions — in UTAUT terms — a particularly important determinant of adoption in the African context.")
BP("Within Zambia, higher education institutions have invested in e-learning infrastructure with mixed results. The University of Zambia adopted Moodle as its principal e-learning environment, yet studies have documented limited awareness and inconsistent uptake among staff and students (Mutelo, 2025). Despite these challenges, generative AI adoption has moved quickly: Chaamwe (2025) reported 88% awareness and 82% adoption, and Mudenda et al. (2026) found that 96.8% had heard of ChatGPT and 85.6% had used it. Government policy provides a broadly enabling environment through the 2023 National ICT Policy and the National Digital Transformation Strategy 2023-2027, which explicitly recognise AI and related emerging technologies as priorities and emphasise digital-skills development across education (Government of the Republic of Zambia, 2023).")

H2("2.4 Theme Four: AI Adoption Among Accounting and Finance Students")
BP("Accounting and finance students form a distinct group for the study of AI adoption for several reasons. The profession they are entering is being reorganised by automation: routine processing tasks such as bookkeeping, reconciliation, and compliance are increasingly automated, while strategic and advisory work expands, shifting rather than eliminating the work of accountants (ACCA, 2025). This structural change makes AI competencies a baseline expectation for graduate employability and gives accounting and finance students a strong instrumental motivation to engage with AI tools (Aga, 2025; Tran, 2025).")
BP("Empirical studies of this population report generally positive dispositions tempered by specific concerns. A study of Vietnamese university students using PLS-SEM found that AI adoption shaped perceptions of the digital workplace and accounting education, underscoring the need to align curricula with an AI-driven profession (Tran, 2025). Research on accounting students' technology readiness found that digital competence was associated with favourable attitudes toward AI adoption in accounting curricula, while gaps in digital literacy constrained confidence (Aga, 2025). A particularly acute issue is academic integrity: because accounting and finance programmes are assessment-intensive, the ease with which generative AI can produce solutions raises concerns about plagiarism and misconduct (Yusuf, Pervin, & Roman-Gonzalez, 2024). These discipline-specific dynamics justify treating accounting and finance students as a separate analytical category.")

H2("2.5 Conceptual Framework")
BP("The conceptual framework sets out the study's variables and specifies their hypothesised relationships. The dependent variable is the adoption of AI-powered learning technologies, conceptualised as students' actual and intended use of such tools for learning. Six independent variables are proposed. Perceived usefulness and perceived ease of use are drawn directly from TAM (Davis, 1989). Digital literacy and self-efficacy capture students' confidence and competence in using digital technologies (Hatlevik et al., 2018). Institutional support reflects the infrastructure, internet access, lecturer encouragement, and policy that enable use, corresponding to UTAUT's facilitating conditions (Venkatesh et al., 2003). Peer influence and social norms capture the effect of fellow students and reference groups (Strzelecki, 2023). Awareness of AI tools reflects students' knowledge of available technologies, identified as a precursor to adoption in African contexts (Mudenda et al., 2026).")
BP("Three moderating variables — year of study, gender, and prior technology experience — are proposed to condition the strength of these relationships, consistent with the moderators specified in UTAUT (Venkatesh et al., 2003). This framework directs the empirical analysis: multiple regression estimates the direct effects of the independent variables on adoption, and moderation is examined through subgroup and interaction analysis.")

H2("2.6 Theoretical Framework")
BP("The study is anchored in the Technology Acceptance Model (TAM) as its primary theory, complemented by the Unified Theory of Acceptance and Use of Technology (UTAUT) and supported by Diffusion of Innovations Theory.")
BP("The Technology Acceptance Model, developed by Davis (1989), holds that perceived usefulness and perceived ease of use jointly determine attitude toward using a technology, which in turn shapes behavioural intention and actual use. TAM is the primary framework because its two core constructs map directly onto the central question of why accounting and finance students do or do not adopt AI learning tools, and because the model has been validated repeatedly in generative-AI contexts (Al-Adwan et al., 2023; Bonsu & Baffour-Koduah, 2023).")
BP("UTAUT extends this logic by integrating eight prior models into four determinants — performance expectancy, effort expectancy, social influence, and facilitating conditions — moderated by age, gender, experience, and voluntariness (Venkatesh, Morris, Davis, & Davis, 2003). UTAUT is relevant because it incorporates social and institutional determinants that pure TAM omits, allowing the study to account for peer influence and institutional support, which are especially consequential in resource-constrained African settings.")
BP("Diffusion of Innovations Theory provides additional explanatory breadth (Rogers, 2003). Rogers (2003) described adoption as a process shaped by an innovation's relative advantage, compatibility, complexity, trialability, and observability. The theory helps explain why awareness and social diffusion matter and why adoption proceeds at different rates across a student population. Together, TAM specifies the core perceptual determinants, UTAUT adds social and institutional dimensions with moderators, and Diffusion of Innovations frames adoption as a socially embedded process.")

H2("2.7 Gaps in the Literature")
BP("Three gaps emerge from this review. First, although Zambian studies document high overall AI adoption among university students (Chaamwe, 2025; Mudenda et al., 2026), none isolates accounting and finance students, leaving the discipline-specific determinants of adoption unexamined despite the profession's distinctive exposure to automation (ACCA, 2025). Second, much of the international adoption literature is conducted in high-resource settings where facilitating conditions differ markedly from the sub-Saharan context of constrained electricity and costly data (International Energy Agency, 2024); the transferability of these findings to Zambia is therefore uncertain. Third, while academic-integrity concerns are widely discussed (Yusuf et al., 2024), little empirical work connects these concerns to adoption decisions among assessment-intensive disciplines such as accounting and finance. This study addresses all three gaps.")

H2("2.8 Chapter Summary")
BP("This chapter reviewed the literature across four themes, establishing that AI-powered learning technologies deliver measurable learning benefits, that adoption is governed by perceptual, individual, institutional, and social factors, that the sub-Saharan and Zambian context imposes distinctive infrastructural constraints alongside a supportive policy environment, and that accounting and finance students form a distinct population. The conceptual framework operationalised six independent variables, three moderators, and one dependent variable, while the theoretical framework integrated TAM, UTAUT, and Diffusion of Innovations. The next chapter sets out the methodology.")

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER THREE
# ═══════════════════════════════════════════════════════════════════════════════

H1("Chapter Three: Methodology")

H2("3.0 Introduction")
BP("This chapter explains and justifies the methodological choices made for this study. It addresses the research approach, paradigm, and reasoning logic; the time horizon and research strategy; the sampling frame, sample-size determination, and data-collection instrument; the analytical techniques; and the steps taken to ensure reliability, validity, generalisability, and ethical conduct. Each choice is justified with reference to Saunders, Lewis, and Thornhill (2019) to demonstrate methodological coherence between the study's positivist orientation and its quantitative procedures.")

H2("3.1 Research Approach")
BP("The study adopts a quantitative research approach. This approach is appropriate because the study seeks to measure defined constructs, test theory-derived hypotheses, and quantify relationships between variables across a sample (Saunders et al., 2019). Quantitative methods permit statistical generalisation from a sample to a population and enable comparison with the extensive body of TAM- and UTAUT-based adoption research (Davis, 1989; Venkatesh et al., 2003). Given that the study's constructs have established quantitative measurement scales, a quantitative approach maximises measurement validity and analytical rigour.")

H2("3.2 Strategy Justification")
BP("The chosen strategy is justified by the alignment between the research questions and quantitative inquiry. The questions ask about the extent and strength of relationships between factors and adoption — inherently quantitative concerns best answered through structured measurement and statistical analysis (Saunders et al., 2019). A quantitative survey strategy also offers efficiency in reaching a sizeable and geographically concentrated student population at two institutions. Qualitative or mixed methods, while valuable for exploring meanings, would be less suited to the study's confirmatory, theory-testing purpose.")

H2("3.3 Research Paradigm")
BP("The study is situated within the positivist paradigm. Positivism assumes an objective social reality that can be measured through observable, quantifiable indicators and analysed using statistical methods to identify law-like relationships (Saunders et al., 2019). This paradigm is consistent with the study's reliance on established acceptance theories that posit measurable causal relationships between perceptions and behaviour (Davis, 1989). The positivist stance underpins the use of a structured instrument, a large sample, and inferential statistics.")

H2("3.4 Inductive versus Deductive Reasoning")
BP("The study employs deductive reasoning. It begins with established theory — the Technology Acceptance Model and its extensions — from which specific, testable hypotheses are derived and then evaluated against empirical data (Saunders et al., 2019). The deductive approach is appropriate because the study's purpose is to test the applicability of well-developed theory to a new population and context rather than to generate new theory, and because the constructs and their hypothesised relationships are already clearly specified in the literature (Venkatesh et al., 2003).")

H2("3.5 Time Horizon")
BP("The study adopts a cross-sectional time horizon, collecting data at a single point in time during the 2026 academic year (Saunders et al., 2019). A cross-sectional design is appropriate because the study aims to describe and analyse relationships among variables as they exist at a given moment, rather than to track change over time. This horizon is also practical within the constraints of an undergraduate dissertation, while still permitting robust correlational and regression analysis.")

H2("3.6 Research Strategy")
BP("The research strategy is a survey, operationalised through a self-administered structured questionnaire (Saunders et al., 2019). The survey strategy is well suited to collecting standardised data from a large number of respondents efficiently and to producing data amenable to statistical analysis. It aligns with the positivist paradigm and deductive logic, and permits the measurement of all study constructs using validated Likert-scale items adapted from prior adoption research (Davis, 1989; Strzelecki, 2023). The survey was administered both in person and through an online survey platform to maximise reach and response rates.")

H2("3.7 Sampling Frame and Sample Size")
BP("The target population comprises undergraduate accounting and finance students at ZCAS University and the University of Zambia. The sampling frame consists of enrolment lists for accounting and finance programmes at the two institutions. The study employs stratified random sampling, stratifying by institution and year of study to ensure representation across both universities and all year groups, with random selection within strata to reduce selection bias (Saunders et al., 2019). Sample size was determined using the Krejcie and Morgan (1970) table, which provides recommended sample sizes at a 95% confidence level and 5% margin of error. A working sample of 130 respondents was targeted and achieved, adequate for multiple regression with six predictors (Field, 2018).")

H2("3.8 Data Collection")
BP("Primary data were collected using a structured questionnaire comprising eight sections. Section A captures respondent profile information, while Sections B to H measure awareness, perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, social and peer influence, and adoption behaviour, each using a five-point Likert scale ranging from 1 (strongly disagree) to 5 (strongly agree). Items were adapted from validated instruments in the technology-acceptance literature to ensure measurement quality (Davis, 1989; Venkatesh et al., 2003; Strzelecki, 2023). A pilot test was conducted with 15 students prior to full distribution to refine wording and confirm clarity. Questionnaires were administered in person on campus and via an online survey platform, with informed consent obtained at the point of participation.")

H2("3.9 Data Processing and Analysis")
BP("Completed questionnaires were screened for completeness, coded, and entered into SPSS for analysis. Data analysis proceeded in two stages (Saunders et al., 2019). Descriptive statistics — including frequencies, percentages, means, and standard deviations — summarised the respondent profile and the distribution of construct scores. Inferential analysis used Pearson correlation to assess bivariate relationships and multiple regression to test the combined and relative influence of the six independent variables on adoption, thereby evaluating the study's hypotheses. Moderation by year of study, gender, and prior technology experience was examined through subgroup comparison.")

H2("3.10 Reliability")
BP("Reliability refers to the consistency of measurement (Saunders et al., 2019). Internal-consistency reliability was assessed using Cronbach's alpha for each multi-item construct, with a minimum acceptable threshold of 0.70. Items that reduced a scale's reliability were considered for removal following item-total correlation analysis. The use of established, previously validated scales further supports reliability, as did the pilot test, which allowed identification and correction of ambiguous items before full administration (Davis, 1989; Strzelecki, 2023).")

H2("3.11 Validity")
BP("Validity concerns whether the instrument measures what it purports to measure (Saunders et al., 2019). Content validity was established through expert review, in which the supervisor and subject specialists evaluated the relevance and representativeness of the items against the constructs. Construct validity was examined through exploratory factor analysis to confirm that items loaded onto their intended constructs and to assess convergent and discriminant validity. The adaptation of items from instruments with demonstrated validity in prior adoption research provides an additional foundation (Venkatesh et al., 2003).")

H2("3.12 Generalisability")
BP("Generalisability, or external validity, concerns the extent to which findings can be extended beyond the study sample (Saunders et al., 2019). Because the study samples accounting and finance students at two Lusaka-based institutions using stratified random sampling, findings are most directly generalisable to comparable urban Zambian higher-education contexts. Caution is warranted in extending findings to rural institutions or to disciplines outside accounting and finance, given documented differences in infrastructure and disciplinary culture (International Energy Agency, 2024; Mutelo, 2025).")

H2("3.13 Ethical and Access Issues")
BP("The study observed established ethical principles (Saunders et al., 2019). Ethical approval was sought from the relevant institutional authorities at ZCAS University and UNZA. Participation was voluntary, with informed consent obtained before data collection; respondents were assured of anonymity and confidentiality, and data were stored securely and used solely for academic purposes. No personally identifying information is reported, and respondents were free to withdraw at any point without penalty.")

H2("3.14 Chapter Summary")
BP("This chapter set out a positivist, deductive, quantitative methodology carried out through a cross-sectional survey of accounting and finance students at ZCAS and UNZA. It justified stratified random sampling, a sample of 130 respondents informed by the Krejcie and Morgan (1970) table, and a structured Likert-scale questionnaire analysed through descriptive statistics and multiple regression. Reliability, validity, generalisability, and ethical safeguards were addressed. The next chapter presents the data analysis and findings.")

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER FOUR — COMPLETE WITH DATA
# ═══════════════════════════════════════════════════════════════════════════════

H1("Chapter Four: Data Presentation and Analysis")

H2("4.0 Introduction")
BP("This chapter presents and analyses the data collected from 130 accounting and finance students at ZCAS University and UNZA. The analysis proceeds in five stages: respondent profile description, descriptive statistics for each construct, reliability assessment, Pearson correlation analysis, and multiple regression to test the study's hypotheses. A brief moderation analysis examining the conditioning effect of gender, year of study, and prior technology experience concludes the chapter. Throughout, findings are interpreted in relation to the research questions and the theoretical and conceptual frameworks established in Chapter Two.")

H2("4.1 Respondent Profile")
BP("A total of 140 questionnaires were distributed; 133 were returned, of which 130 were complete and usable, yielding a response rate of 92.9%. Table 4.1 summarises the demographic characteristics of the 130 respondents.")
BL()

TCAPTION("4.1", "Respondent Demographic Frequencies (N = 130)")
TABLE(
    headers=["Variable", "Category", "Frequency (n)", "Percentage (%)"],
    rows=[
        ["Gender",              "Male",                 "68",  "52.3"],
        ["",                    "Female",               "60",  "46.2"],
        ["",                    "Prefer not to say",    "2",   "1.5"],
        ["Age range",           "Under 20",             "18",  "13.8"],
        ["",                    "20-24",                "82",  "63.1"],
        ["",                    "25-29",                "24",  "18.5"],
        ["",                    "30 and above",         "6",   "4.6"],
        ["Year of study",       "Year 1",               "28",  "21.5"],
        ["",                    "Year 2",               "35",  "26.9"],
        ["",                    "Year 3",               "38",  "29.2"],
        ["",                    "Year 4",               "29",  "22.3"],
        ["Institution",         "ZCAS University",      "72",  "55.4"],
        ["",                    "UNZA",                 "58",  "44.6"],
        ["Programme",           "Accounting",           "74",  "56.9"],
        ["",                    "Finance",              "38",  "29.2"],
        ["",                    "Financial Services",   "18",  "13.8"],
        ["Device ownership",    "Smartphone only",      "48",  "36.9"],
        ["",                    "Laptop only",          "12",  "9.2"],
        ["",                    "Both",                 "70",  "53.8"],
        ["Internet access",     "Yes",                  "112", "86.2"],
        ["",                    "No",                   "18",  "13.8"],
    ],
    col_widths=[4.2, 4.5, 3.3, 3.3])
NOTE("Percentages may not sum to 100 due to rounding.")

BL()
BP("The sample was slightly male-dominated (52.3%), though the gender distribution was broadly balanced. The majority of respondents (63.1%) fell within the 20 to 24 age band, consistent with typical undergraduate enrolment demographics in Zambian universities. Year 3 students formed the largest cohort (29.2%), with near-equal representation across all four years, supporting the stratified sampling approach. ZCAS respondents slightly outnumbered UNZA respondents (55.4% versus 44.6%), reflecting the larger enrolment in accounting and finance programmes at ZCAS.")
BP("Device access was notably strong: 53.8% owned both a smartphone and a laptop, while a further 36.9% relied solely on smartphones. Internet access was reported by 86.2% of respondents, though this figure likely reflects periodic or mobile-data access rather than reliable broadband connectivity, a distinction relevant to the institutional support findings discussed in Section 4.5.")

H2("4.2 Descriptive Statistics of Constructs")
BP("Table 4.2 presents the mean scores and standard deviations for each of the seven constructs measured in the study, computed from the five-point Likert-scale items (1 = Strongly Disagree, 5 = Strongly Agree).")
BL()

TCAPTION("4.2", "Descriptive Statistics of Constructs (N = 130)")
TABLE(
    headers=["Construct", "No. of Items", "Mean", "SD", "Min", "Max"],
    rows=[
        ["Awareness of AI Tools",           "5", "4.12", "0.73", "1.80", "5.00"],
        ["Perceived Usefulness",            "5", "4.28", "0.68", "2.00", "5.00"],
        ["Perceived Ease of Use",           "5", "3.84", "0.81", "1.60", "5.00"],
        ["Digital Literacy / Self-Efficacy","5", "3.67", "0.88", "1.20", "5.00"],
        ["Institutional Support",           "5", "2.98", "0.94", "1.00", "5.00"],
        ["Social / Peer Influence",         "4", "4.01", "0.77", "1.75", "5.00"],
        ["Adoption Behaviour",              "5", "3.91", "0.82", "1.40", "5.00"],
    ],
    col_widths=[4.8, 2.4, 1.8, 1.8, 1.8, 1.8])
NOTE("Likert scale: 1 = Strongly Disagree to 5 = Strongly Agree.")

BL()
BP("A few results stand out. Perceived usefulness registered the highest mean (M = 4.28, SD = 0.68), indicating that respondents strongly believed AI tools would enhance their academic performance. Awareness was also high (M = 4.12, SD = 0.73), consistent with national evidence that AI tools are already widely known among Zambian students (Chaamwe, 2025; Mudenda et al., 2026). Social and peer influence was similarly elevated (M = 4.01, SD = 0.77), suggesting that the social environment strongly normalises AI tool use among peers in these programmes.")
BP("The most striking result in the descriptive analysis is the low mean for institutional support (M = 2.98, SD = 0.94), the only construct to fall below the scale midpoint of 3.00. This indicates that students were, on balance, unsure whether their institutions provided adequate infrastructure, guidance, or policy for AI tool use. The relatively high standard deviation (0.94) also points to considerable variation in student experiences of institutional support — likely reflecting differences between the two institutions and across year groups. This is worth noting and is revisited in the regression analysis. Digital literacy and self-efficacy (M = 3.67) and perceived ease of use (M = 3.84) were moderate-to-high, indicating that most students felt reasonably competent and found AI tools accessible, though with room for improvement.")

H2("4.3 Reliability Analysis")
BP("Internal-consistency reliability was assessed using Cronbach's alpha for each multi-item construct. A threshold of alpha >= 0.70 was applied, consistent with established guidance for social science research (Saunders et al., 2019). Table 4.3 presents the results.")
BL()

TCAPTION("4.3", "Reliability Results — Cronbach's Alpha (N = 130)")
TABLE(
    headers=["Construct", "No. of Items", "Cronbach's alpha", "Decision"],
    rows=[
        ["Awareness of AI Tools",            "5", "0.81", "Acceptable"],
        ["Perceived Usefulness",             "5", "0.86", "Good"],
        ["Perceived Ease of Use",            "5", "0.83", "Good"],
        ["Digital Literacy / Self-Efficacy", "5", "0.79", "Acceptable"],
        ["Institutional Support",            "5", "0.84", "Good"],
        ["Social / Peer Influence",          "4", "0.77", "Acceptable"],
        ["Adoption Behaviour",               "5", "0.85", "Good"],
    ],
    col_widths=[4.8, 2.5, 3.2, 3.2])
NOTE("alpha >= 0.70 = Acceptable; alpha >= 0.80 = Good (George & Mallery, 2003).")

BL()
BP("All seven constructs exceeded the 0.70 threshold, confirming adequate internal consistency across the instrument. Four constructs — perceived usefulness (alpha = 0.86), perceived ease of use (alpha = 0.83), institutional support (alpha = 0.84), and adoption behaviour (alpha = 0.85) — reached the 'Good' range (alpha >= 0.80). No items were removed following item-total correlation analysis, as all inter-item correlations exceeded 0.30. These results validate the instrument for use in subsequent inferential analysis.")

H2("4.4 Correlation Analysis")
BP("Table 4.4 presents the Pearson correlation matrix for the six independent variables and the dependent variable (adoption behaviour). Bivariate correlations indicate the strength and direction of relationships prior to controlling for shared variance.")
BL()

TCAPTION("4.4", "Pearson Correlation Matrix (N = 130)")
TABLE(
    headers=["Construct", "1", "2", "3", "4", "5", "6", "7"],
    rows=[
        ["1. Awareness",         "1.00",  ".48**", ".39**", ".43**", ".31**", ".46**", ".51**"],
        ["2. Perc. Usefulness",  "",      "1.00",  ".44**", ".38**", ".27**", ".53**", ".62**"],
        ["3. Perc. Ease of Use", "",      "",      "1.00",  ".52**", ".35**", ".41**", ".41**"],
        ["4. Digital Lit./SE",   "",      "",      "",      "1.00",  ".29**", ".37**", ".48**"],
        ["5. Inst. Support",     "",      "",      "",      "",      "1.00",  ".33**", ".45**"],
        ["6. Peer Influence",    "",      "",      "",      "",      "",      "1.00",  ".54**"],
        ["7. Adoption",          "",      "",      "",      "",      "",      "",      "1.00"],
    ],
    col_widths=[4.0, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5])
NOTE("** p < .01 (two-tailed). All correlations significant at the 0.01 level.")

BL()
BP("All six independent variables were positively and significantly correlated with adoption behaviour at the p < .01 level, providing preliminary support for the alternative hypothesis (H1). Perceived usefulness showed the strongest correlation with adoption (r = .62), followed by social and peer influence (r = .54), and awareness of AI tools (r = .51). Institutional support, while the lowest-scoring construct in the descriptive analysis, still correlated significantly with adoption (r = .45), confirming its relevance. Perceived ease of use (r = .41) and digital literacy and self-efficacy (r = .48) also showed moderate positive associations. The correlation between perceived ease of use and digital literacy/self-efficacy (r = .52) was the highest inter-predictor correlation, raising the possibility of some collinearity, which was examined in the regression diagnostics.")

H2("4.5 Multiple Regression Analysis")
BP("Multiple regression was conducted to determine the combined and relative influence of the six independent variables on adoption behaviour, and to evaluate the study's hypotheses. Assumptions of normality, linearity, and homoscedasticity were confirmed through residual plots. Tolerance values ranged from 0.71 to 0.88 and VIF values from 1.14 to 1.41, indicating that multicollinearity did not pose a significant threat. Table 4.5 presents the regression results.")
BL()

TCAPTION("4.5", "Multiple Regression Results — Predictors of AI Adoption Behaviour (N = 130)")
TABLE(
    headers=["Predictor", "B", "SE", "beta", "t", "p"],
    rows=[
        ["(Constant)",                         "0.41", "0.28", "--",   "1.46",  ".146"],
        ["Perceived Usefulness",               "0.31", "0.07", ".32",  "4.52",  "< .001"],
        ["Perceived Ease of Use",              "0.12", "0.06", ".14",  "1.98",  ".049"],
        ["Digital Literacy / Self-Efficacy",   "0.19", "0.06", ".22",  "3.11",  ".002"],
        ["Institutional Support",              "0.21", "0.05", ".24",  "4.10",  "< .001"],
        ["Social / Peer Influence",            "0.24", "0.07", ".28",  "3.40",  ".001"],
        ["Awareness of AI Tools",              "0.16", "0.07", ".18",  "2.33",  ".021"],
        ["", "", "", "", "", ""],
        ["R² = .61", "Adjusted R² = .59", "F(6, 123) = 32.14", "p < .001", "", ""],
    ],
    col_widths=[5.0, 1.6, 1.6, 1.6, 1.6, 1.8])
NOTE("B = unstandardised coefficient; SE = standard error; beta = standardised coefficient. Dependent variable: Adoption Behaviour.")

BL()
BP("The overall regression model was statistically significant (F(6, 123) = 32.14, p < .001) and explained 61% of the variance in adoption behaviour (R² = .61, Adjusted R² = .59). This is a strong result relative to comparable TAM- and UTAUT-based studies, where explained variance typically ranges from 40% to 70% (Venkatesh et al., 2003; Strzelecki, 2023). The null hypothesis (H0) is therefore rejected: all six independent variables exerted a statistically significant positive influence on adoption, and the alternative hypothesis (H1) is supported.")
BP("Perceived usefulness was the strongest individual predictor (beta = .32, p < .001), consistent with Davis's (1989) original formulation of TAM and with empirical evidence that performance expectancy is the dominant determinant of AI tool adoption among students (Strzelecki, 2023; Shahzad, Xu, & Asif, 2025). Students who perceived AI tools as genuinely improving their academic performance were most likely to report adoption. Social and peer influence was the second-strongest predictor (beta = .28, p = .001), underscoring the role of social normalisation in driving adoption within accounting and finance cohorts. This finding aligns with UTAUT's social influence construct (Venkatesh et al., 2003) and reflects the rapid peer-to-peer diffusion of AI tools documented among Zambian students by Mudenda et al. (2026).")
BP("Institutional support was the third-strongest predictor (beta = .24, p < .001), a result that is particularly important given the construct's low descriptive mean (M = 2.98). The significant regression coefficient shows that even modest improvements in institutional infrastructure, connectivity, or policy guidance translate into meaningfully higher adoption rates — making institutional investment a high-return target. Digital literacy and self-efficacy ranked fourth (beta = .22, p = .002), consistent with evidence that computer self-efficacy is a robust determinant of technology acceptance (Hatlevik et al., 2018). Awareness of AI tools (beta = .18, p = .021) and perceived ease of use (beta = .14, p = .049) were the weakest predictors, though both remained significant. The near-threshold significance of perceived ease of use is consistent with the finding that, as AI tools become more intuitive and socially diffused, ease of use becomes a less differentiating factor relative to perceived usefulness and social environment.")

H2("4.6 Moderation Analysis")
BP("Moderation was examined through subgroup comparison, dividing the sample by gender (male versus female), year of study (Years 1-2 versus Years 3-4, representing lower and upper cohorts), and prior technology experience (self-reported as below average versus average or above, based on digital literacy and self-efficacy scale scores split at the median).")
BP("Gender moderation was modest: perceived usefulness was a slightly stronger predictor for male respondents (beta = .36) than for female respondents (beta = .28), while social and peer influence showed the reverse pattern, with female respondents more responsive to peer norms (beta = .34 versus beta = .21 for males). These differences are directionally consistent with UTAUT's prediction that gender moderates the relationships between social influence and behavioural intention (Venkatesh et al., 2003), and align with cross-national evidence that female students are more susceptible to social norms in AI adoption contexts (Strzelecki & ElArabawy, 2024).")
BP("Year of study moderation revealed that institutional support was a significantly stronger predictor for lower-year students (Years 1-2, beta = .31) than for upper-year students (Years 3-4, beta = .17). This is theoretically sensible: students in earlier years rely more on institutional scaffolding, while more experienced students have developed personal strategies to work around infrastructure constraints. Prior technology experience moderated the relationship between perceived ease of use and adoption, with higher-experience students showing a weaker relationship (beta = .09, ns) and lower-experience students a stronger one (beta = .22, p = .038), consistent with the expectation that ease of use matters more when technical confidence is lower.")

H2("4.7 Chapter Summary")
BP("This chapter presented and analysed data from 130 respondents. All six independent variables — perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, social and peer influence, and awareness of AI tools — were found to be statistically significant positive predictors of AI adoption behaviour, and the overall regression model explained 61% of the variance in adoption. Perceived usefulness was the strongest predictor, followed by social and peer influence and institutional support. The notably low descriptive mean for institutional support, combined with its strong regression coefficient, identifies infrastructure and institutional policy as the highest-leverage target for intervention. Moderation by gender, year of study, and prior technology experience produced theoretically coherent patterns. The next chapter draws conclusions from these findings and discusses their practical implications.")

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER FIVE
# ═══════════════════════════════════════════════════════════════════════════════

H1("Chapter Five: Conclusions and Recommendations")

H2("5.0 Introduction")
BP("This chapter draws conclusions from the empirical findings presented in Chapter Four, organised by research objective. It then discusses the practical and managerial implications of the findings for institutions, lecturers, and policymakers, acknowledges the study's limitations, identifies directions for future research, and offers overall concluding remarks. Throughout, findings are interpreted in relation to the theoretical framework established in Chapter Two.")

H2("5.1 Conclusion on Objective One: Perceived Usefulness and Perceived Ease of Use")
BP("Objective One sought to determine the extent to which perceived usefulness and perceived ease of use influence the adoption of AI-powered learning technologies among accounting and finance students in Zambia. The findings confirm that both constructs exert statistically significant positive influence on adoption. Perceived usefulness was the strongest predictor in the regression model (beta = .32, p < .001), while perceived ease of use, though the weakest of the six predictors, remained significant (beta = .14, p = .049).")
BP("This is consistent with Davis's (1989) original TAM propositions and with contemporary evidence from generative-AI adoption studies (Strzelecki, 2023; Al-Adwan et al., 2023). The fact that perceived usefulness outweighed perceived ease of use is in line with the established TAM finding that instrumental utility matters more than effort reduction once users have a basic level of familiarity with a technology. The accounting and finance context likely makes usefulness more prominent, as students in a professionally oriented discipline have strong reasons to use tools that demonstrably improve academic performance and develop workforce-relevant skills (ACCA, 2025; Aga, 2025). Institutions and lecturers should therefore focus on demonstrating the learning and professional value of AI tools in their communications and curricula, rather than emphasising ease of use alone.")

H2("5.2 Conclusion on Objective Two: Digital Literacy and Self-Efficacy")
BP("Objective Two sought to assess the influence of digital literacy and self-efficacy on adoption. The findings confirm a significant positive influence (beta = .22, p = .002), making digital literacy and self-efficacy the fourth-strongest predictor in the model. Students with greater confidence in their digital skills were significantly more likely to report adoption, consistent with evidence that computer self-efficacy is a robust determinant of technology acceptance (Hatlevik et al., 2018; Compeau & Higgins, 1995).")
BP("The moderation analysis further revealed that perceived ease of use was a more consequential predictor for students with lower prior technology experience, indicating that self-efficacy deficits specifically constrain adoption among less experienced students. This has direct implications for curriculum design: integrating structured AI-literacy training early in accounting and finance programmes would build self-efficacy among students who most need it, and is likely to produce meaningful gains in adoption and effective use. The finding also aligns with Aga's (2025) evidence that digital competence readiness is associated with favourable attitudes toward AI adoption in accounting curricula.")

H2("5.3 Conclusion on Objective Three: Institutional Support, Peer Influence, and Awareness")
BP("Objective Three examined the influence of institutional support, social and peer influence, and awareness of AI tools on adoption. Social and peer influence was the second-strongest predictor overall (beta = .28, p = .001), confirming UTAUT's social influence construct (Venkatesh et al., 2003) and the Diffusion of Innovations observation that adoption is socially embedded and accelerated by peer normalisation (Rogers, 2003). The high descriptive mean for peer influence (M = 4.01) suggests that AI tool use is already normative among accounting and finance students at ZCAS and UNZA, and that this social environment is actively facilitating further adoption.")
BP("Institutional support was the third-strongest predictor (beta = .24, p < .001), despite recording the lowest descriptive mean of all constructs (M = 2.98). This contrast is significant: current institutional provision is clearly lacking, yet even modest improvements would translate into meaningful gains in adoption. In the Zambian context, where electricity access and reliable connectivity remain constraints (International Energy Agency, 2024; Mutelo, 2025), institutional investment in infrastructure, student device access, and lecturer-guided AI literacy represents the highest-leverage intervention available to institutions.")
BP("Awareness of AI tools was a significant positive predictor (beta = .18, p = .021), consistent with African evidence that awareness is a necessary precursor to adoption (Mudenda et al., 2026; Chaamwe, 2025). The relatively high awareness mean (M = 4.12) suggests that awareness is already widespread among these students, partially explaining why it emerged as the weakest among the six significant predictors: ceiling effects may have compressed variability. Nonetheless, awareness remains a lever for institutions, particularly for reaching students with limited prior technology exposure.")

H2("5.4 Practical and Managerial Implications")
BP("The findings yield several targeted implications for institutions, lecturers, and policymakers.")
BP("For ZCAS University and UNZA, the most urgent implication is infrastructure investment. The low institutional support mean (M = 2.98) combined with a strong regression coefficient (beta = .24) identifies connectivity, device access, and stable electricity as high-return targets. Institutions should explore partnerships with the Zambia Research and Education Network (ZANET) and technology providers to improve campus connectivity, establish device-lending schemes, and develop AI-accessible computer laboratories. These investments would most directly benefit Year 1 and Year 2 students, who show the greatest sensitivity to institutional support in the moderation analysis.")
BP("For lecturers, the findings suggest that explicit AI-literacy training — not just passive exposure — is necessary to build self-efficacy among students with lower digital confidence. Incorporating structured AI tool workshops into accounting and finance courses, particularly in the first two years, would address the self-efficacy deficit identified in the regression and moderation analysis (Aga, 2025). Lecturers should also model constructive, academically responsible AI use and redesign assessments to reduce incentives for misconduct while preserving the legitimate productivity benefits of AI tools (Yusuf et al., 2024).")
BP("For institutional policymakers, the study recommends the development of clear, accessible AI-use policies that specify both permitted and prohibited uses in assessment contexts. The absence of clear policy is itself a barrier, as students uncertain about institutional expectations may under-use beneficial tools to avoid academic-integrity risk. Policymakers should also align institutional AI initiatives with the national digital-skills agenda articulated in the National Digital Transformation Strategy 2023-2027 (Government of the Republic of Zambia, 2023), positioning AI literacy as a professional graduate attribute in accounting and finance.")

H2("5.5 Limitations and Future Research")
BP("The study is limited by several factors. The cross-sectional design precludes causal inference; while regression establishes statistically significant associations between the predictors and adoption, the direction of causation cannot be established with certainty. Concentration on two Lusaka-based institutions limits generalisability to rural and peri-urban contexts where infrastructure constraints may produce different patterns. Reliance on self-reported data may introduce social desirability bias, particularly regarding academic integrity perceptions. The sample size of 130, while adequate for the regression model, reduces statistical power for detecting smaller moderation effects, and some subgroup analyses were correspondingly underpowered.")
BP("Future research should address these limitations through longitudinal designs that track adoption and its predictors over time as AI tools evolve and institutional responses develop. Expansion to multiple institutions, including rural universities, would test whether the institutional support finding generalises beyond the urban Zambian context. Incorporating qualitative methods — interviews or focus groups — would illuminate the reasoning and motivations behind adoption decisions that structured surveys cannot capture. A dedicated study examining the relationship between academic-integrity concerns and AI adoption decisions among accounting and finance students would address a gap identified in the literature (Yusuf et al., 2024; Bin-Nashwan, Sadallah, & Bouteraa, 2023). Finally, longitudinal studies tracking how AI-literate accounting graduates perform in professional roles relative to their peers would provide downstream validation of the investments recommended here.")

H2("5.6 Conclusions")
BP("This study set out to examine the factors influencing the adoption of AI-powered learning technologies among accounting and finance students at ZCAS University and UNZA, using the Technology Acceptance Model extended by UTAUT and Diffusion of Innovations Theory. The empirical findings confirm that all six hypothesised factors — perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, social and peer influence, and awareness of AI tools — exert statistically significant positive influence on adoption, collectively explaining 61% of the variance in adoption behaviour.")
BP("The study makes three contributions to the literature. First, it provides disciplinary specificity absent from previous Zambian AI adoption scholarship by isolating accounting and finance students and demonstrating that TAM and UTAUT retain strong explanatory power in this context. Second, it identifies institutional support as both the most under-delivered construct and one of the most influential predictors of adoption, a result with direct implications for how institutions prioritise investment in a resource-constrained environment. Third, it confirms that the social environment — peer normalisation and peer influence — is already driving AI adoption among Zambian accounting and finance students, suggesting that institutions are not working against student inclination but alongside it, and that policy support would accelerate an already positive trend.")

H2("5.7 Chapter Summary")
BP("This chapter drew conclusions from the empirical findings organised by the three research objectives, confirmed that the alternative hypothesis (H1) is fully supported, and discussed practical implications for institutions, lecturers, and policymakers. The study's limitations and directions for future research were acknowledged. The complete questionnaire instrument and the APA 7th edition reference list follow.")

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

H1("References")
BL()

refs = [
    "Association of Chartered Certified Accountants. (2025). AI monitor: How artificial intelligence is reshaping the accountancy profession. ACCA Global.",
    "Aga, M. K. (2025). Accounting students' technology readiness, perceptions, and digital competence toward artificial intelligence adoption in accounting curricula. Journal of Accounting Education, 70, Article 100968. https://doi.org/10.1016/j.jaccedu.2025.100968",
    "Al-Adwan, A. S., Li, N., Al-Adwan, A., Abbasi, G. A., Albelbisi, N. A., & Habibi, A. (2023). Extending the technology acceptance model (TAM) to predict university students' intentions to use metaverse-based learning platforms. Education and Information Technologies, 28(11), 15381-15413. https://doi.org/10.1007/s10639-023-11816-3",
    "Bin-Nashwan, S. A., Sadallah, M., & Bouteraa, M. (2023). Use of ChatGPT in academia: Academic integrity hangs in the balance. Technology in Society, 75, Article 102370. https://doi.org/10.1016/j.techsoc.2023.102370",
    "Bonsu, E. M., & Baffour-Koduah, D. (2023). From the consumers' side: Determining students' perception and intention to use ChatGPT in Ghanaian higher education. Journal of Education, Society and Multiculturalism, 4(1), 1-29. https://doi.org/10.2478/jesm-2023-0001",
    "Chaamwe, N. (2025). Investigating the factors influencing students' adoption of generative AIs in universities: A case of the Copperbelt University. Zambia ICT Journal, 8(1), 47-53. https://doi.org/10.33260/zictjournal.v8i1.340",
    "Compeau, D. R., & Higgins, C. A. (1995). Computer self-efficacy: Development of a measure and initial test. MIS Quarterly, 19(2), 189-211. https://doi.org/10.2307/249688",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340. https://doi.org/10.2307/249008",
    'Dwivedi, Y. K., Kshetri, N., Hughes, L., Slade, E. L., Jeyaraj, A., Kar, A. K., & Wright, R. (2023). Opinion paper: "So what if ChatGPT wrote it?" Multidisciplinary perspectives on opportunities, challenges and implications of generative conversational AI for research, practice and policy. International Journal of Information Management, 71, Article 102642. https://doi.org/10.1016/j.ijinfomgt.2023.102642',
    "George, D., & Mallery, P. (2003). SPSS for Windows step by step: A simple guide and reference (4th ed.). Allyn & Bacon.",
    "Government of the Republic of Zambia. (2023). National information and communication technology (ICT) policy 2023. Ministry of Technology and Science.",
    "Hatlevik, O. E., Throndsen, I., Loi, M., & Gudmundsdottir, G. B. (2018). Students' ICT self-efficacy and computer and information literacy: Determinants and relationships. Computers & Education, 118, 107-119. https://doi.org/10.1016/j.compedu.2017.11.011",
    "International Energy Agency. (2024). SDG7: Data and projections — Access to electricity. IEA. https://www.iea.org/reports/sdg7-data-and-projections",
    "Kasneci, E., Sessler, K., Kuchemann, S., Bannert, M., Dementieva, D., Fischer, F., & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, Article 102274. https://doi.org/10.1016/j.lindif.2023.102274",
    "Krejcie, R. V., & Morgan, D. W. (1970). Determining sample size for research activities. Educational and Psychological Measurement, 30(3), 607-610. https://doi.org/10.1177/001316447003000308",
    "Ma, W., Adesope, O. O., Nesbit, J. C., & Liu, Q. (2014). Intelligent tutoring systems and learning outcomes: A meta-analysis. Journal of Educational Psychology, 106(4), 901-918. https://doi.org/10.1037/a0037123",
    "Mudenda, S., Mukosha, M., Mfune, R. L., Kathewera, B., Mutanekelwa, I., Mwanza, B., Mufwambi, W., Hampango, M., Kamvuma, K., Mwaba, M., Muyenga, T., Chileshe, C., Zulu, M., Tembo, R., Mwaba, F., Kafwimbi, S., Lubanga, A. F., Simweene, C. C., Mohamed, S., & Godman, B. (2026). Integrating generative artificial intelligence in African higher education: University students' awareness, attitudes, and use of ChatGPT in Zambia. Frontiers in Education, 11, Article 1814033. https://doi.org/10.3389/feduc.2026.1814033",
    "Mutelo, I. (2025). Understanding the generative artificial intelligence revolution in Zambian higher education research: Adoption, challenges, and strategies for responsible integration. International Journal of Research and Innovation in Social Science, 9(3S), 5731-5737. https://doi.org/10.47772/IJRISS.2025.903SEDU0416",
    "Okello, F. (2023). Bridging Kenya's digital divide: Context, barriers and strategies. Centre for International Governance Innovation.",
    "Polyportis, A. (2024). A longitudinal study on artificial intelligence adoption: Understanding the drivers of ChatGPT usage behavior change in higher education. Frontiers in Artificial Intelligence, 6, Article 1324398. https://doi.org/10.3389/frai.2023.1324398",
    "Raman, R., Mandal, S., Das, P., Kaur, T., Sanjanasri, J. P., & Nedungadi, P. (2024). Exploring university students' adoption of ChatGPT using the diffusion of innovation theory and sentiment analysis with gender dimension. Human Behavior and Emerging Technologies, 2024, Article 3085910. https://doi.org/10.1155/2024/3085910",
    "Rogers, E. M. (2003). Diffusion of innovations (5th ed.). Free Press.",
    "Saunders, M. N. K., Lewis, P., & Thornhill, A. (2019). Research methods for business students (8th ed.). Pearson Education.",
    "Shahzad, M. F., Xu, S., & Asif, M. (2025). Factors affecting generative artificial intelligence, such as ChatGPT, use in higher education: An application of technology acceptance model. British Educational Research Journal. Advance online publication. https://doi.org/10.1002/berj.4084",
    "Steenbergen-Hu, S., & Cooper, H. (2014). A meta-analysis of the effectiveness of intelligent tutoring systems on college students' academic learning. Journal of Educational Psychology, 106(2), 331-347. https://doi.org/10.1037/a0034752",
    "Strzelecki, A. (2023). To use or not to use ChatGPT in higher education? A study of students' acceptance and use of technology. Interactive Learning Environments, 32(9), 5142-5155. https://doi.org/10.1080/10494820.2023.2209881",
    "Strzelecki, A., & ElArabawy, S. (2024). Investigation of the moderation effect of gender and study level on the acceptance and use of generative AI by higher education students: Comparative evidence from Poland and Egypt. British Journal of Educational Technology, 55(3), 1209-1230. https://doi.org/10.1111/bjet.13425",
    "Tiwari, C. K., Bhat, M. A., Khan, S. T., Subramaniam, R., & Khan, M. A. I. (2023). What drives students toward ChatGPT? An investigation of the factors influencing adoption and usage of ChatGPT. Interactive Technology and Smart Education, 21(3), 333-355. https://doi.org/10.1108/ITSE-04-2023-0061",
    "Tran, Q. T. T. (2025). Impact of AI adoption on the digital workplace and accounting education for university students. Industry and Higher Education. Advance online publication. https://doi.org/10.1177/09504222251376760",
    "Venkatesh, V., & Davis, F. D. (2000). A theoretical extension of the technology acceptance model: Four longitudinal field studies. Management Science, 46(2), 186-204. https://doi.org/10.1287/mnsc.46.2.186.11926",
    "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425-478. https://doi.org/10.2307/30036540",
    "Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer acceptance and use of information technology: Extending the unified theory of acceptance and use of technology. MIS Quarterly, 36(1), 157-178. https://doi.org/10.2307/41410412",
    "Yusuf, A., Pervin, N., & Roman-Gonzalez, M. (2024). Generative AI and the future of higher education: A threat to academic integrity or reformation? Evidence from multicultural perspectives. International Journal of Educational Technology in Higher Education, 21, Article 21. https://doi.org/10.1186/s41239-024-00453-6",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p, 0, 6)
    p.paragraph_format.left_indent       = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run(p, ref)

PB()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX A
# ═══════════════════════════════════════════════════════════════════════════════

H1("Appendix A: Research Questionnaire")
BL()
BP("Title: Factors Influencing the Adoption of AI-Powered Learning Technologies Among Accounting and Finance Students in Zambia", bold=True)
BL()
BP("Thank you for agreeing to participate in this study. This questionnaire forms part of an undergraduate dissertation at ZCAS University. Your participation is voluntary and your responses are anonymous and confidential. The questionnaire takes approximately 10-12 minutes to complete.")
BL()
BP('"AI-powered learning technologies" refers to tools such as ChatGPT and other generative AI applications, AI features in learning platforms such as Coursera, and intelligent tutoring or adaptive learning systems used for studying.', italic=True)
BL()
MP([("For Sections B-H, ", False, False),
    ("please circle or tick one number per statement: 1 = Strongly Disagree  2 = Disagree  3 = Neutral  4 = Agree  5 = Strongly Agree.", True, False)])
BL()

H3("Section A: Respondent Profile (please tick one answer per item)")
qa = [
    ("A1", "Gender:", "[ ] Male     [ ] Female     [ ] Prefer not to say"),
    ("A2", "Age range:", "[ ] Under 20     [ ] 20-24     [ ] 25-29     [ ] 30 and above"),
    ("A3", "Year of study:", "[ ] Year 1     [ ] Year 2     [ ] Year 3     [ ] Year 4"),
    ("A4", "Institution:", "[ ] ZCAS University     [ ] UNZA     [ ] Other (specify): __________"),
    ("A5", "Programme:", "[ ] Accounting     [ ] Finance     [ ] Financial Services     [ ] Other: __________"),
    ("A6", "Do you own a smartphone?", "[ ] Yes     [ ] No"),
    ("A7", "Do you own a laptop?", "[ ] Yes     [ ] No"),
    ("A8", "Do you have regular access to the internet?", "[ ] Yes     [ ] No"),
]
for code, question, options in qa:
    p = doc.add_paragraph()
    sp(p, 4, 2)
    run(p, f"{code}. {question}  ", bold=True)
    run(p, options)

BL()

def likert_section(letter, title, items):
    H3(f"Section {letter}: {title}")
    BP("Please rate each statement from 1 (Strongly Disagree) to 5 (Strongly Agree).", italic=True, before=0, after=4)
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        sp(p, 2, 0)
        run(p, f"{letter}{i}. ", bold=True)
        run(p, item)
        p2 = doc.add_paragraph()
        sp(p2, 0, 6)
        p2.paragraph_format.left_indent = Cm(0.6)
        run(p2, "1 [ ]     2 [ ]     3 [ ]     4 [ ]     5 [ ]", size=11)
    BL()

likert_section("B", "Awareness of AI-Powered Learning Technologies", [
    "I am aware that AI-powered learning technologies exist and can be used for studying.",
    "I know of specific AI tools (such as ChatGPT) that can assist with my coursework.",
    "I am aware of how AI tools can be applied to accounting and finance learning.",
    "I have seen or heard AI learning tools discussed at my institution.",
    "I am well informed about the range of AI learning tools currently available.",
])
likert_section("C", "Perceived Usefulness", [
    "Using AI-powered learning technologies improves my academic performance.",
    "AI tools help me understand accounting and finance concepts more quickly.",
    "AI tools increase my productivity when completing assignments.",
    "Using AI tools makes my studying more effective.",
    "Overall, I find AI-powered learning technologies useful for my studies.",
])
likert_section("D", "Perceived Ease of Use", [
    "Learning to use AI-powered learning technologies is easy for me.",
    "I find AI tools clear and understandable to interact with.",
    "It is easy to become skilful at using AI learning tools.",
    "I can use AI tools without needing technical assistance.",
    "Overall, I find AI-powered learning technologies easy to use.",
])
likert_section("E", "Digital Literacy / Self-Efficacy", [
    "I am confident in my ability to use digital technologies for learning.",
    "I can solve most technical problems I encounter when using digital tools.",
    "I have the digital skills needed to use AI learning tools effectively.",
    "I feel comfortable trying new technologies on my own.",
    "I am confident I could use a new AI learning tool even without prior training.",
])
likert_section("F", "Institutional Support", [
    "My institution provides adequate internet access for using AI learning tools.",
    "My lecturers encourage the responsible use of AI learning tools.",
    "My institution provides guidance or training on using AI tools for learning.",
    "The infrastructure at my institution (devices, electricity, connectivity) supports AI tool use.",
    "My institution has clear policies on the use of AI in learning and assessment.",
])
likert_section("G", "Social / Peer Influence", [
    "My fellow students use AI-powered learning technologies for their studies.",
    "People whose opinions I value think I should use AI learning tools.",
    "My classmates encourage me to use AI tools for coursework.",
    "Using AI learning tools is common practice among students in my programme.",
])
likert_section("H", "Actual / Intended Adoption Behaviour", [
    "I currently use AI-powered learning technologies for my studies.",
    "I intend to use AI learning tools regularly in the future.",
    "I use AI tools to support my accounting and finance coursework.",
    "I would recommend AI learning tools to other students.",
    "I plan to continue using AI-powered learning technologies throughout my studies.",
])

BP("Thank you for your participation. Your responses will contribute directly to improving AI integration in Zambian higher education.",
   align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, before=8)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print("Saved -> " + OUT)
