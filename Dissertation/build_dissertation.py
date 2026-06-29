"""
Builds the fully formatted Booklesss dissertation DOCX.
Dikhilani Mvula — Student No. 202101786
ZCAS University / UNZA — Bachelor of Arts in Financial Services
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"c:\Users\deeky\OneDrive\Desktop\Booklesss\Dissertation\Mvula_Dissertation_Final.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_spacing(para, before=0, after=0, line=None, line_rule=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line
    if line_rule:
        pf.line_spacing_rule = line_rule

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def body_para(doc, text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              size=12, before=0, after=6, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=18, after=6)
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=12, after=4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=8, after=4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, italic=True)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_blank(doc, n=1):
    for _ in range(n):
        body_para(doc, before=0, after=0)

def mixed_para(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, indent=None):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before=before, after=after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def bullet_para(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=3)
    run = p.add_run(text)
    set_font(run)
    return p

def numbered_para(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=3)
    run = p.add_run(text)
    set_font(run)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def section_label(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=10, after=2)
    run = p.add_run(label)
    set_font(run, bold=True, size=12)
    return p


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins: 2.54 cm all sides (standard academic)
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.18)   # left wider for binding
    section.right_margin  = Cm(2.54)
    section.page_width    = Cm(21.0)   # A4
    section.page_height   = Cm(29.7)

# Default paragraph font
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)
doc.styles["Normal"].paragraph_format.line_spacing = Pt(24)  # double-space
doc.styles["Normal"].paragraph_format.space_after  = Pt(0)


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

add_blank(doc, 3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=6)
run = p.add_run("AN EXAMINATION OF FACTORS INFLUENCING THE ADOPTION OF\nAI-POWERED LEARNING TECHNOLOGIES AMONG ACCOUNTING AND\nFINANCE STUDENTS IN SELECTED HIGHER LEARNING INSTITUTIONS\nIN ZAMBIA")
set_font(run, size=14, bold=True)

add_blank(doc, 2)

for line in [
    "By",
    "",
    "DIKHILANI MVULA",
    "Student Number: 202101786",
]:
    p = body_para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)

add_blank(doc)

body_para(doc, "A Dissertation Submitted in Partial Fulfilment of the Requirements for the Award of the Degree of Bachelor of Arts in Financial Services",
          align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)

add_blank(doc)

for line in [
    "ZCAS University in Association with the University of Zambia",
    "",
    "Department of Accounting and Finance",
    "",
    "Supervisor Code: NM",
    "",
    "Lusaka, Zambia",
    "",
    "2026",
]:
    body_para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "DECLARATION")
add_blank(doc)

body_para(doc, "I, Dikhilani Mvula (Student Number: 202101786), declare that this dissertation is my own original work and has not been submitted for any other degree or award at this or any other institution. All sources consulted have been duly acknowledged in the text and in the reference list.")
add_blank(doc)
body_para(doc, "Where the work of others has been used, this has been properly cited and referenced in accordance with APA 7th edition conventions.")
add_blank(doc, 3)

body_para(doc, "Signature: _________________________________")
add_blank(doc)
body_para(doc, "Date: ____________________________________")
add_blank(doc)
body_para(doc, "Name: Dikhilani Mvula")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "ABSTRACT")
add_blank(doc)

body_para(doc,
    "Artificial intelligence (AI)-powered learning technologies are now widely used among Zambian university students, yet the specific determinants of adoption among accounting and finance students remain unexamined. This study investigates the factors influencing the adoption of AI-powered learning technologies among accounting and finance students at ZCAS University and the University of Zambia (UNZA). Grounded in the Technology Acceptance Model (Davis, 1989) and extended by the Unified Theory of Acceptance and Use of Technology (Venkatesh, Morris, Davis, & Davis, 2003) and Diffusion of Innovations Theory (Rogers, 2003), the study examines six independent variables — perceived usefulness, perceived ease of use, digital literacy and self-efficacy, institutional support, social and peer influence, and awareness of AI tools — together with three moderating variables: year of study, gender, and prior technology experience.")
add_blank(doc)
body_para(doc,
    "A positivist, deductive, quantitative methodology is adopted. Data will be collected through a structured questionnaire administered to a stratified random sample of approximately 100 to 150 undergraduate accounting and finance students, with analysis conducted using descriptive statistics and multiple regression. The study isolates a population not previously examined in isolation in the Zambian literature, contributing disciplinary specificity to a field in which existing evidence treats university students as a single, undifferentiated group. Findings are anticipated to demonstrate that perceived usefulness, peer influence, and institutional support are significant positive determinants of adoption, with infrastructure remaining a binding constraint in the sub-Saharan context. Results will inform institutional policy, AI-literacy curriculum design, and academic-integrity frameworks for accounting and finance programmes in Zambia.")
add_blank(doc)
mixed_para(doc, [("Keywords: ", True, False),
                 ("technology acceptance model, AI adoption, accounting education, higher education, Zambia, generative AI", False, True)])

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "ACKNOWLEDGEMENTS")
add_blank(doc)

body_para(doc,
    "I am deeply grateful to my dissertation supervisor (Supervisor Code: NM) for expert guidance, constructive feedback, and continued encouragement throughout the research process. The direction and insight provided have been invaluable in shaping the rigour and focus of this study.")
add_blank(doc)
body_para(doc,
    "I also wish to acknowledge the academic staff and administration at both ZCAS University and the University of Zambia for facilitating access to institutional resources and research participants. My gratitude extends to my family and fellow students for their unwavering support and motivation.")
add_blank(doc)
body_para(doc,
    "This research is dedicated to all Zambian students working toward a future in which technology and education are aligned with national development aspirations.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "TABLE OF CONTENTS")
add_blank(doc)

toc_entries = [
    ("Declaration", "ii"),
    ("Abstract", "iii"),
    ("Acknowledgements", "iv"),
    ("Table of Contents", "v"),
    ("List of Tables", "vi"),
    ("List of Abbreviations", "vii"),
    ("", ""),
    ("CHAPTER ONE: INTRODUCTION", ""),
    ("1.0  Introduction", "1"),
    ("1.1  Background of the Study", "1"),
    ("1.2  Research Problem", "2"),
    ("1.3  Justification", "3"),
    ("1.4  Research Aim", "3"),
    ("1.5  Research Objectives", "4"),
    ("1.6  Research Questions", "4"),
    ("1.7  Research Hypothesis", "4"),
    ("1.8  Research Scope", "5"),
    ("1.9  Research Contributions", "5"),
    ("1.10 Research Design", "6"),
    ("1.11 Research Approach and Method", "6"),
    ("1.12 Data Collection and Analysis Techniques", "7"),
    ("1.13 Dissertation Layout", "7"),
    ("1.14 Chapter Summary", "8"),
    ("", ""),
    ("CHAPTER TWO: LITERATURE REVIEW", ""),
    ("2.0  Introduction", "9"),
    ("2.1  AI-Powered Learning Technologies in Higher Education", "9"),
    ("2.2  Factors Influencing Technology Adoption", "10"),
    ("2.3  Technology Adoption in Sub-Saharan Africa and the Zambian Context", "11"),
    ("2.4  AI Adoption Among Accounting and Finance Students", "12"),
    ("2.5  Conceptual Framework", "13"),
    ("2.6  Theoretical Framework", "14"),
    ("2.7  Gaps in the Literature", "15"),
    ("2.8  Chapter Summary", "15"),
    ("", ""),
    ("CHAPTER THREE: METHODOLOGY", ""),
    ("3.0  Introduction", "16"),
    ("3.1  Research Approach", "16"),
    ("3.2  Strategy Justification", "17"),
    ("3.3  Research Paradigm", "17"),
    ("3.4  Inductive versus Deductive", "17"),
    ("3.5  Time Horizon", "18"),
    ("3.6  Research Strategy", "18"),
    ("3.7  Sampling Frame and Sample Size", "18"),
    ("3.8  Data Collection", "19"),
    ("3.9  Data Processing and Analysis", "20"),
    ("3.10 Reliability", "20"),
    ("3.11 Validity", "21"),
    ("3.12 Generalisability", "21"),
    ("3.13 Ethical and Access Issues", "22"),
    ("3.14 Chapter Summary", "22"),
    ("", ""),
    ("CHAPTER FOUR: DATA ANALYSIS PLAN AND PRESENTATION", ""),
    ("4.0  Introduction", "23"),
    ("4.1  Analysis Plan", "23"),
    ("4.2  Template Table Shells", "24"),
    ("", ""),
    ("CHAPTER FIVE: CONCLUSIONS FRAMEWORK", ""),
    ("5.0  Introduction", "28"),
    ("5.1  Conclusion on Objective One", "28"),
    ("5.2  Conclusion on Objective Two", "29"),
    ("5.3  Conclusion on Objective Three", "29"),
    ("5.4  Conclusion on Objective Four", "30"),
    ("5.5  Practical and Managerial Implications", "30"),
    ("5.6  Limitations and Future Research", "31"),
    ("5.7  Conclusions", "31"),
    ("5.8  Summary", "32"),
    ("", ""),
    ("References", "33"),
    ("Appendix A: Research Questionnaire", "38"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    if not entry:
        set_spacing(p, before=0, after=4)
        continue
    is_chapter = entry.startswith("CHAPTER")
    set_spacing(p, before=4 if is_chapter else 0, after=2)
    tab = p.add_run(entry)
    set_font(tab, bold=is_chapter)
    if page:
        p.add_run("\t")
        pg = p.add_run(page)
        set_font(pg, bold=is_chapter)
    # Right-align page numbers using tab stop
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    pPr = p._p.get_or_add_pPr()
    tabs = _OE("w:tabs")
    tab_el = _OE("w:tab")
    tab_el.set(_qn("w:val"), "right")
    tab_el.set(_qn("w:leader"), "dot")
    tab_el.set(_qn("w:pos"), "9072")  # ~16cm
    tabs.append(tab_el)
    pPr.append(tabs)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "LIST OF TABLES")
add_blank(doc)

tables_list = [
    ("Table 4.1", "Respondent Demographic Frequencies", "24"),
    ("Table 4.2", "Descriptive Statistics of Constructs", "25"),
    ("Table 4.3", "Reliability Results (Cronbach's Alpha)", "26"),
    ("Table 4.4", "Correlation Matrix", "26"),
    ("Table 4.5", "Multiple Regression Results", "27"),
]

for num, title, page in tables_list:
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=4)
    r1 = p.add_run(f"{num}: {title}")
    set_font(r1)
    p.add_run("\t")
    r2 = p.add_run(page)
    set_font(r2)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_el = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:leader"), "dot")
    tab_el.set(qn("w:pos"), "9072")
    tabs.append(tab_el)
    pPr.append(tabs)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "LIST OF ABBREVIATIONS")
add_blank(doc)

abbrevs = [
    ("ACCA", "Association of Chartered Certified Accountants"),
    ("AI", "Artificial Intelligence"),
    ("APA", "American Psychological Association"),
    ("ChatGPT", "Chat Generative Pre-trained Transformer"),
    ("IEA", "International Energy Agency"),
    ("ITS", "Intelligent Tutoring System"),
    ("PLS-SEM", "Partial Least Squares Structural Equation Modelling"),
    ("TAM", "Technology Acceptance Model"),
    ("UNZA", "University of Zambia"),
    ("UTAUT", "Unified Theory of Acceptance and Use of Technology"),
    ("ZCAS", "Zambia Centre for Accountancy Studies"),
    ("ZANET", "Zambia Research and Education Network"),
    ("ZMW", "Zambian Kwacha"),
]

for abbr, meaning in abbrevs:
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=3)
    r1 = p.add_run(f"{abbr}")
    set_font(r1, bold=True)
    r2 = p.add_run(f"\t{meaning}")
    set_font(r2)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_el = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "left")
    tab_el.set(qn("w:pos"), "1800")
    tabs.append(tab_el)
    pPr.append(tabs)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER ONE
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "CHAPTER ONE: INTRODUCTION")

heading2(doc, "1.0 Introduction")
body_para(doc,
    "This chapter establishes the foundation of a study that examines the factors influencing the adoption of artificial intelligence (AI)-powered learning technologies among accounting and finance students in selected higher learning institutions in Zambia. The rapid diffusion of generative AI applications, adaptive learning systems, and intelligent tutoring platforms has altered how students in higher education access information, complete assessments, and develop professional competencies (Strzelecki, 2023). Within accounting and finance, where data processing and analytical reasoning are central, these technologies carry particular significance because the profession itself is being reorganised around automation and advisory work (ACCA, 2025). The chapter presents the background to the study, articulates the research problem, and justifies why an investigation grounded in the Zambian context is warranted. It then states the aim, objectives, questions, and hypotheses that direct the inquiry, before outlining the scope, anticipated contributions, and the methodological choices that structure the remaining chapters. In doing so, the chapter situates the study within established theories of technology acceptance while signalling the distinctive empirical contribution it intends to make to Zambian higher education scholarship.")

heading2(doc, "1.1 Background of the Study")
body_para(doc,
    "Artificial intelligence has moved from a peripheral curiosity to a central feature of teaching and learning in higher education within a remarkably short period. The public release of generative AI applications, most prominently ChatGPT in late 2022, accelerated student engagement with conversational AI for writing, problem-solving, and revision (Strzelecki, 2023). Globally, empirical studies confirm that students perceive these tools as useful and easy to use, and that such perceptions translate into strong behavioural intention to adopt them (Bonsu & Baffour-Koduah, 2023; Al-Adwan, Li, Al-Adwan, Abbasi, Albelbisi, & Habibi, 2023). Alongside generative AI, adaptive learning systems and intelligent tutoring systems have demonstrated consistent learning gains across educational levels and subject domains (Ma, Adesope, Nesbit, & Liu, 2014).")
body_para(doc,
    "In Zambia, adoption has been swift despite infrastructural limitations. Chaamwe (2025) reported that 88% of surveyed university students were aware of generative AI and 82% had adopted it for learning, while Mudenda et al. (2026), in a multi-institutional study, found that 96.8% of students had heard of ChatGPT and 85.6% had used it. These figures suggest that AI-powered learning technologies are already embedded in student practice. However, adoption is uneven and shaped by digital infrastructure, awareness, digital literacy, and institutional readiness (Mutelo, 2025). Accounting and finance students occupy a distinctive position in this development because their discipline is undergoing structural transformation through automation, machine learning, and data analytics, which heightens both the relevance of AI competencies and the urgency of understanding how these students engage with AI learning tools (ACCA, 2025; Aga, 2025).")

heading2(doc, "1.2 Research Problem")
body_para(doc,
    "Although generative AI use among Zambian university students is now widespread (Chaamwe, 2025; Mudenda et al., 2026), the existing evidence treats the student body as a single, undifferentiated population. No located Zambian study isolates accounting and finance students, despite the fact that this group faces discipline-specific pressures arising from the automation of routine accounting tasks and the growing demand for technology-literate finance professionals (ACCA, 2025; Tran, 2025). Consequently, the factors that specifically drive or inhibit AI adoption among accounting and finance students in Zambia remain poorly understood. The problem is compounded by structural constraints that characterise sub-Saharan higher education, including limited electricity access, high data costs, and uneven device ownership (International Energy Agency, 2024; Mutelo, 2025). Without disciplinary and contextual specificity, institutions risk designing AI integration policies that neither reflect the realities of Zambian accounting and finance students nor address the academic-integrity concerns that AI introduces into assessment-intensive programmes (Yusuf, Pervin, & Román-González, 2024). This study addresses that gap by examining, through an established theoretical framework, the determinants of AI adoption among accounting and finance students at ZCAS University and the University of Zambia.")

heading2(doc, "1.3 Justification")
body_para(doc,
    "This study is justified on theoretical, practical, and contextual grounds. Theoretically, it applies and extends the Technology Acceptance Model (Davis, 1989) to a generative AI context and a population that has not been examined in isolation in Zambia, thereby testing the model's continued explanatory relevance for emerging technologies (Al-Adwan et al., 2023). Practically, the findings will assist universities, lecturers, and policymakers in designing targeted interventions, such as AI-literacy training and assessment redesign, that respond to the actual determinants of adoption rather than to assumptions (Mudenda et al., 2026). Contextually, the study responds to national priorities articulated in Zambia's digital transformation agenda, which identifies digital skills and emerging technologies, including AI, as central to economic development (Government of the Republic of Zambia, 2023). By concentrating on accounting and finance students, the study also speaks directly to the future of a profession in which AI competencies are becoming a baseline expectation for graduate employability (ACCA, 2025; Aga, 2025). The investigation therefore contributes evidence that is at once theoretically grounded, practically actionable, and contextually anchored in Zambian higher education.")

heading2(doc, "1.4 Research Aim")
body_para(doc,
    "The aim of this study is to examine the factors influencing the adoption of AI-powered learning technologies among accounting and finance students in selected higher learning institutions in Zambia, with a view to informing institutional policy and pedagogical practice.")

heading2(doc, "1.5 Research Objectives")
body_para(doc, "The study is guided by four objectives:")
numbered_para(doc, "To determine the extent to which perceived usefulness and perceived ease of use influence the adoption of AI-powered learning technologies among accounting and finance students in Zambia.")
numbered_para(doc, "To assess the influence of individual factors, specifically digital literacy and self-efficacy, on the adoption of AI-powered learning technologies among these students.")
numbered_para(doc, "To examine the influence of institutional support and social or peer influence on the adoption of AI-powered learning technologies among these students.")
numbered_para(doc, "To evaluate the extent to which awareness of AI tools, together with the moderating effects of year of study, gender, and prior technology experience, shapes adoption among accounting and finance students.")

heading2(doc, "1.6 Research Questions")
body_para(doc, "Corresponding to the objectives, the study addresses four questions:")
numbered_para(doc, "To what extent do perceived usefulness and perceived ease of use influence the adoption of AI-powered learning technologies among accounting and finance students in Zambia?")
numbered_para(doc, "How do digital literacy and self-efficacy influence the adoption of AI-powered learning technologies among these students?")
numbered_para(doc, "What is the influence of institutional support and social or peer influence on the adoption of AI-powered learning technologies among these students?")
numbered_para(doc, "To what extent does awareness of AI tools, together with year of study, gender, and prior technology experience, shape adoption among these students?")

heading2(doc, "1.7 Research Hypothesis")
body_para(doc, "The study tests the following hypotheses, which will be evaluated through multiple regression analysis:")
add_blank(doc)
mixed_para(doc, [("H0 (Null hypothesis): ", True, False), ("Perceived usefulness, perceived ease of use, digital literacy/self-efficacy, institutional support, peer influence, and awareness of AI tools have no statistically significant influence on the adoption of AI-powered learning technologies among accounting and finance students in Zambia.", False, False)])
add_blank(doc)
mixed_para(doc, [("H1 (Alternative hypothesis): ", True, False), ("Perceived usefulness, perceived ease of use, digital literacy/self-efficacy, institutional support, peer influence, and awareness of AI tools have a statistically significant positive influence on the adoption of AI-powered learning technologies among accounting and finance students in Zambia.", False, False)])

heading2(doc, "1.8 Research Scope")
body_para(doc,
    "The study is confined to undergraduate accounting and finance students at two institutions, ZCAS University and the University of Zambia (UNZA), both located in Lusaka. The geographical scope is therefore Lusaka-based higher education, while the conceptual scope is limited to factors associated with the adoption of AI-powered learning technologies as conceptualised within the Technology Acceptance Model and its extensions (Davis, 1989; Venkatesh, Morris, Davis, & Davis, 2003). The temporal scope is cross-sectional, capturing student perceptions and reported behaviour at a single point during the 2026 academic year. The study examines learning-oriented AI technologies, including generative AI applications such as ChatGPT, AI features in platforms such as Coursera, and intelligent tutoring or adaptive learning systems, rather than enterprise accounting software or professional audit automation tools. These boundaries ensure analytical focus while acknowledging that findings are most directly transferable to comparable urban Zambian institutions.")

heading2(doc, "1.9 Research Contributions")
body_para(doc,
    "The study makes three contributions. First, it extends the empirical literature on AI adoption in Zambian higher education by isolating accounting and finance students, a population not previously examined separately (Chaamwe, 2025; Mudenda et al., 2026). Second, it contributes theoretically by testing the explanatory power of the Technology Acceptance Model for generative AI within a sub-Saharan, discipline-specific context, thereby informing debates about the model's continued relevance for emerging technologies (Al-Adwan et al., 2023). Third, it offers practical contributions by generating evidence that institutions can use to design AI-literacy initiatives, infrastructure investments, and academic-integrity policies tailored to accounting and finance programmes (ACCA, 2025; Yusuf et al., 2024). Collectively, these contributions respond to a recognised gap in contextual and disciplinary specificity and align with Zambia's national emphasis on digital skills development (Government of the Republic of Zambia, 2023).")

heading2(doc, "1.10 Research Design")
body_para(doc,
    "The study adopts a descriptive-correlational design, which is appropriate for describing the prevalence of AI adoption and the strength and direction of relationships between the identified factors and adoption behaviour (Saunders, Lewis, & Thornhill, 2019). A descriptive component captures the distribution of awareness, perceptions, and reported adoption among respondents, while a correlational component quantifies the associations between the independent variables and the dependent variable. This design suits a study seeking to test theory-derived hypotheses about relationships rather than to establish causation through experimentation. It also aligns with the positivist orientation of the research and supports the use of structured quantitative instruments and inferential statistical techniques (Saunders et al., 2019).")

heading2(doc, "1.11 Research Approach and Method")
body_para(doc,
    "The study employs a deductive approach, beginning with established theory — the Technology Acceptance Model and its extensions — from which testable hypotheses are derived and then evaluated against empirical data (Saunders et al., 2019). The method is quantitative, using a structured questionnaire to collect numerical data amenable to statistical analysis. This combination of a deductive logic and a quantitative method is consistent with the positivist paradigm and with the study's objective of measuring relationships between defined variables across a sample of students (Davis, 1989; Venkatesh et al., 2003). The quantitative method also enables the study to produce findings that can be compared with the substantial body of TAM- and UTAUT-based adoption research conducted internationally and in African contexts (Strzelecki, 2023; Mudenda et al., 2026).")

heading2(doc, "1.12 Data Collection and Analysis Techniques")
body_para(doc,
    "Primary data will be collected through a self-administered, structured questionnaire using a five-point Likert scale, distributed to a stratified random sample of accounting and finance students at ZCAS University and UNZA. The instrument adapts validated items from prior TAM and UTAUT studies to ensure measurement quality (Davis, 1989; Venkatesh et al., 2003; Strzelecki, 2023). Data analysis will proceed in two stages. Descriptive statistics, including frequencies, means, and standard deviations, will summarise respondent profiles and construct scores. Inferential analysis will employ Pearson correlation and multiple regression to test the hypothesised relationships between the independent variables and adoption (Saunders et al., 2019). Reliability will be assessed using Cronbach's alpha, with a threshold of 0.70, and validity will be established through expert content review and exploratory factor analysis. Data will be analysed using statistical software to ensure accuracy and reproducibility.")

heading2(doc, "1.13 Dissertation Layout")
body_para(doc,
    "The dissertation is organised into five chapters. Chapter One introduces the study, presenting the background, problem, objectives, and methodological orientation. Chapter Two reviews the literature across four themes, develops the conceptual framework, and elaborates the theoretical framework. Chapter Three details the methodology, justifying each methodological choice with reference to established research-methods scholarship. Chapter Four, presented in outline form, sets out the data-analysis plan and template tables to be populated once primary data are collected. Chapter Five presents conclusions organised by research objective, discusses implications and limitations, and identifies directions for future research. A complete questionnaire and an APA-formatted reference list are appended.")

heading2(doc, "1.14 Chapter Summary")
body_para(doc,
    "This chapter has introduced the study and established its rationale. It has shown that while AI-powered learning technologies are now widely used among Zambian university students, the determinants of adoption among accounting and finance students specifically remain unexamined. The chapter articulated the aim, four objectives, four questions, and the hypotheses to be tested, and it outlined a positivist, deductive, quantitative design employing a structured questionnaire and regression analysis. The next chapter reviews the relevant literature and develops the conceptual and theoretical frameworks that guide the empirical investigation.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER TWO
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "CHAPTER TWO: LITERATURE REVIEW")

heading2(doc, "2.0 Introduction")
body_para(doc,
    "This chapter reviews the scholarly literature relevant to the adoption of AI-powered learning technologies among accounting and finance students, organising the discussion around four themes. The first theme examines what AI-powered learning technologies are and the evidence of their impact in higher education. The second considers the factors that influence technology adoption, drawing on dominant acceptance models. The third situates the discussion within the sub-Saharan African and Zambian context, attending to infrastructure, policy, and barriers. The fourth focuses specifically on accounting and finance students as a distinct population. The chapter then presents the conceptual framework that operationalises the study's variables and the theoretical framework that underpins the analysis, before identifying the gaps the study addresses and summarising the review. Throughout, the chapter draws on peer-reviewed scholarship, prioritising recent empirical work while grounding the discussion in foundational theory.")

heading2(doc, "2.1 Theme One: AI-Powered Learning Technologies in Higher Education")
body_para(doc,
    "AI-powered learning technologies encompass a range of applications that use machine learning, natural language processing, and adaptive algorithms to support teaching and learning. These include intelligent tutoring systems that model a learner's knowledge state and deliver individualised instruction, adaptive learning platforms that adjust content to learner pace, and generative AI applications such as ChatGPT that produce conversational, on-demand assistance (Strzelecki, 2023). Platforms such as Coursera have similarly integrated AI-driven features to personalise learning pathways. The defining characteristic of these technologies is adaptivity: they respond to individual learner characteristics, needs, and progress (Ma, Adesope, Nesbit, & Liu, 2014).")
body_para(doc,
    "The evidence base on learning impact is substantial for the more established forms of these technologies. A meta-analysis synthesising 107 effect sizes from 14,321 participants found that the use of intelligent tutoring systems was associated with greater achievement than teacher-led, large-group instruction (g = .42), non-ITS computer-based instruction (g = .57), and textbooks or workbooks (g = .35) (Ma et al., 2014). Positive effects were observed across nearly all subject domains and educational levels, indicating broad applicability. Generative AI applications are newer, and the evidence on their learning outcomes is still emerging, but early studies report that students perceive them as enhancing engagement, comprehension, and efficiency (Bonsu & Baffour-Koduah, 2023; Polyportis, 2024). In accounting and finance education specifically, AI tools are being applied to automated grading, finance simulations, and AI-assisted tutoring, although the scholarship on their measured effectiveness in this discipline remains in early development (Aga, 2025; Tran, 2025). This unevenness — robust evidence for established tools and emerging evidence for generative AI — motivates continued empirical attention.")

heading2(doc, "2.2 Theme Two: Factors Influencing Technology Adoption")
body_para(doc,
    "The study of technology adoption is dominated by two complementary models. The Technology Acceptance Model proposes that two beliefs — perceived usefulness and perceived ease of use — are the fundamental determinants of a user's intention to adopt a technology and subsequent usage behaviour (Davis, 1989). Perceived usefulness refers to the degree to which a person believes a technology will enhance performance, while perceived ease of use refers to the belief that using it will be effortless. The Unified Theory of Acceptance and Use of Technology synthesised eight prior models into four core determinants, namely performance expectancy, effort expectancy, social influence, and facilitating conditions, moderated by age, gender, experience, and voluntariness (Venkatesh, Morris, Davis, & Davis, 2003). UTAUT explained up to 70% of the variance in behavioural intention, a substantial improvement over earlier individual models.")
body_para(doc,
    "Empirical applications to AI tools confirm the relevance of these constructs while highlighting context-dependence. In a PLS-SEM study of 534 Polish university students, habit was the strongest predictor of behavioural intention to use ChatGPT, followed by performance expectancy and hedonic motivation, with social influence and effort expectancy exerting smaller effects (Strzelecki, 2023). A study of 411 U.S. university students found that perceived usefulness and subjective norm significantly predicted intention to use ChatGPT, with perceived ease of use and trust operating indirectly through usefulness, with the model explaining roughly 51% of the variance in usage intention. Beyond these model constructs, individual factors such as digital literacy, prior technology experience, and self-efficacy shape adoption. Computer self-efficacy — defined as confidence in one's ability to use technology — is a consistent predictor of technology acceptance and effective use (Hatlevik, Throndsen, Loi, & Gudmundsdóttir, 2018). Institutional factors, including infrastructure, internet access, lecturer support, and policy, and social factors such as peer influence and subjective norms, further condition whether favourable perceptions translate into actual adoption (Strzelecki, 2023; Mudenda et al., 2026).")

heading2(doc, "2.3 Theme Three: Technology Adoption in Sub-Saharan Africa and the Zambian Context")
body_para(doc,
    "Technology adoption in sub-Saharan Africa unfolds against a backdrop of significant infrastructural constraint. The International Energy Agency (2024) reported that around 600 million people in sub-Saharan Africa — approximately 47% of the population — lacked access to electricity, with the region accounting for roughly 80% of the global population without such access. High data costs, limited 4G coverage, and uneven device ownership compound this electricity deficit, producing a persistent digital divide that disproportionately affects rural and low-income learners (Okello, 2023). These conditions establish facilitating conditions, in UTAUT terms, as a particularly salient determinant of adoption in the African context.")
body_para(doc,
    "Within Zambia, higher education institutions have invested in e-learning infrastructure with mixed results. The University of Zambia adopted Moodle as its principal e-learning environment, supported by connectivity through the Zambia Research and Education Network, yet studies have documented limited awareness and inconsistent uptake among staff and students (Mutelo, 2025). Despite these constraints, generative AI adoption has been rapid. Chaamwe (2025) reported 88% awareness and 82% adoption among Zambian university students, and Mudenda et al. (2026) found that 96.8% had heard of ChatGPT and 85.6% had used it. Government policy provides an enabling environment: the SMART Zambia transformation agenda was launched in 2015, with subsequent instruments including the 2023 National ICT Policy and the National Digital Transformation Strategy 2023–2027 explicitly recognising AI, the internet of things, and related emerging technologies as priorities and emphasising digital-skills development across education (Government of the Republic of Zambia, 2023). The policy environment is therefore broadly supportive, even as implementation lags behind ambition (Mutelo, 2025).")

heading2(doc, "2.4 Theme Four: AI Adoption Among Accounting and Finance Students")
body_para(doc,
    "Accounting and finance students constitute a distinct population for the study of AI adoption for several reasons. The profession they are entering is being reorganised by automation: routine processing tasks such as bookkeeping, reconciliation, and compliance are increasingly automated, while strategic and advisory work expands, shifting rather than eliminating the work of accountants (ACCA, 2025). This structural change makes AI competencies a baseline expectation for graduate employability and gives accounting and finance students a strong instrumental motivation to engage with AI tools (Aga, 2025; Tran, 2025).")
body_para(doc,
    "Empirical studies of this population report generally positive dispositions tempered by specific concerns. A study of Vietnamese university students using PLS-SEM found that AI adoption shaped perceptions of the digital workplace and accounting education, underscoring the need to align curricula and infrastructure with an AI-driven profession (Tran, 2025). Research on accounting students' technology readiness found that digital competence and readiness were associated with favourable attitudes toward AI adoption in accounting curricula, while gaps in digital literacy constrained confidence (Aga, 2025). Attitudes among business and finance students may differ from those in other disciplines because of the field's data-intensive, rules-based character and its professional-accreditation requirements. A particularly acute issue is academic integrity. Because accounting and finance programmes are assessment-intensive, the ease with which generative AI can produce solutions raises concerns about plagiarism and misconduct, a tension documented across higher education and especially relevant where assessments determine professional eligibility (Yusuf, Pervin, & Román-González, 2024). These discipline-specific dynamics justify treating accounting and finance students as a separate analytical category rather than subsuming them within the general student body.")

heading2(doc, "2.5 Conceptual Framework")
body_para(doc,
    "The conceptual framework operationalises the study's variables and specifies their hypothesised relationships. The dependent variable is the adoption of AI-powered learning technologies, conceptualised as students' actual and intended use of such tools for learning. Six independent variables are proposed. Perceived usefulness and perceived ease of use are drawn directly from the Technology Acceptance Model and represent the belief that AI tools enhance learning performance and require little effort, respectively (Davis, 1989). Digital literacy and self-efficacy capture students' confidence and competence in using digital technologies (Hatlevik et al., 2018). Institutional support reflects the infrastructure, internet access, lecturer encouragement, and policy that enable use, corresponding to UTAUT's facilitating conditions (Venkatesh et al., 2003). Peer influence and social norms capture the effect of fellow students and reference groups (Strzelecki, 2023). Awareness of AI tools reflects students' knowledge of available technologies, identified as a precursor to adoption in African contexts (Mudenda et al., 2026).")
body_para(doc,
    "The framework posits that each independent variable exerts a positive influence on adoption. Three moderating variables — year of study, gender, and prior technology experience — are proposed to condition the strength of these relationships, consistent with the moderators specified in UTAUT (Venkatesh et al., 2003). For example, prior technology experience is expected to strengthen the relationship between perceived ease of use and adoption, while gender and year of study may moderate the influence of social and institutional factors. This framework directs the empirical analysis: multiple regression will estimate the direct effects of the independent variables on adoption, and moderation will be examined through subgroup or interaction analysis.")

heading2(doc, "2.6 Theoretical Framework")
body_para(doc,
    "The study is anchored in the Technology Acceptance Model (TAM) as its primary theory, complemented by the Unified Theory of Acceptance and Use of Technology (UTAUT) and supported by Diffusion of Innovations Theory.")
body_para(doc,
    "The Technology Acceptance Model, developed by Davis (1989), holds that perceived usefulness and perceived ease of use jointly determine attitude toward using a technology, which in turn shapes behavioural intention and actual use. Davis (1989) demonstrated that perceived usefulness correlated strongly with usage and that perceived ease of use influenced usefulness, establishing a parsimonious and highly replicable model. TAM is the primary framework for this study because its two core constructs map directly onto the central question of why accounting and finance students do or do not adopt AI learning tools, and because the model has been validated repeatedly in generative-AI contexts (Al-Adwan et al., 2023; Bonsu & Baffour-Koduah, 2023).")
body_para(doc,
    "UTAUT extends this logic by integrating eight prior models into four determinants — performance expectancy, effort expectancy, social influence, and facilitating conditions — moderated by age, gender, experience, and voluntariness (Venkatesh, Morris, Davis, & Davis, 2003). UTAUT is relevant here because it incorporates social and institutional determinants that pure TAM omits, allowing the study to account for peer influence and institutional support, which are especially consequential in resource-constrained African settings (Strzelecki, 2023; Mudenda et al., 2026). Its moderators also justify the study's moderating variables.")
body_para(doc,
    "Diffusion of Innovations Theory provides supporting explanatory breadth (Rogers, 2003). Rogers (2003) characterised adoption as a process influenced by an innovation's relative advantage, compatibility, complexity, trialability, and observability, and classified adopters from innovators to laggards. The theory illuminates why awareness and social diffusion matter and why adoption proceeds at different rates across a student population. Together, these three theories furnish a coherent lens: TAM specifies the core perceptual determinants, UTAUT adds social and institutional dimensions with moderators, and Diffusion of Innovations frames adoption as a socially embedded process.")

heading2(doc, "2.7 Gaps in the Literature")
body_para(doc,
    "Three gaps emerge from this review. First, although Zambian studies document high overall AI adoption among university students (Chaamwe, 2025; Mudenda et al., 2026), none isolates accounting and finance students, leaving the discipline-specific determinants of adoption unexamined despite the profession's distinctive exposure to automation (ACCA, 2025). Second, much of the international adoption literature is conducted in high-resource settings where facilitating conditions differ markedly from the sub-Saharan context of constrained electricity and costly data (International Energy Agency, 2024; Strzelecki, 2023); the transferability of these findings to Zambia is therefore uncertain. Third, while academic-integrity concerns are widely discussed in general terms (Yusuf et al., 2024), little empirical work connects these concerns to adoption decisions among assessment-intensive disciplines such as accounting and finance. This study addresses all three gaps by examining the determinants of AI adoption among accounting and finance students at ZCAS and UNZA using an established theoretical framework adapted to the Zambian context.")

heading2(doc, "2.8 Chapter Summary")
body_para(doc,
    "This chapter reviewed the literature across four themes, establishing that AI-powered learning technologies deliver measurable learning benefits, that adoption is governed by perceptual, individual, institutional, and social factors, that the sub-Saharan and Zambian context imposes distinctive infrastructural constraints alongside a supportive policy environment, and that accounting and finance students form a distinct population owing to the automation of their profession. The conceptual framework operationalised six independent variables, three moderators, and one dependent variable, while the theoretical framework integrated TAM, UTAUT, and Diffusion of Innovations. The identified gaps justify the empirical study. The next chapter sets out the methodology through which these relationships will be tested.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER THREE
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "CHAPTER THREE: METHODOLOGY")

heading2(doc, "3.0 Introduction")
body_para(doc,
    "This chapter explains and justifies the methodological choices that structure the empirical investigation. It addresses the research approach, paradigm, and reasoning logic; the time horizon and research strategy; the sampling frame, sample-size determination, and data-collection instrument; the analytical techniques; and the steps taken to ensure reliability, validity, generalisability, and ethical conduct. Each choice is justified with reference to established research-methods scholarship, principally Saunders, Lewis, and Thornhill (2019), to demonstrate methodological coherence between the study's positivist orientation and its quantitative procedures.")

heading2(doc, "3.1 Research Approach")
body_para(doc,
    "The study adopts a quantitative research approach. This approach is appropriate because the study seeks to measure defined constructs, test theory-derived hypotheses, and quantify relationships between variables across a sample (Saunders, Lewis, & Thornhill, 2019). Quantitative methods permit statistical generalisation from a sample to a population and enable comparison with the extensive body of TAM- and UTAUT-based adoption research (Davis, 1989; Venkatesh, Morris, Davis, & Davis, 2003). Given that the study's constructs, such as perceived usefulness and self-efficacy, have established quantitative measurement scales, a quantitative approach maximises measurement validity and analytical rigour.")

heading2(doc, "3.2 Strategy Justification")
body_para(doc,
    "The chosen strategy is justified by the alignment between the research questions and quantitative inquiry. The questions ask about the extent and strength of relationships between factors and adoption, which are inherently quantitative concerns best answered through structured measurement and statistical analysis (Saunders et al., 2019). A quantitative survey strategy also offers efficiency in reaching a sizeable and geographically concentrated student population at two institutions, and it supports the deductive testing of the study's hypotheses. Qualitative or mixed methods, while valuable for exploring meanings, would be less suited to the study's confirmatory, theory-testing purpose.")

heading2(doc, "3.3 Research Paradigm")
body_para(doc,
    "The study is situated within the positivist paradigm. Positivism assumes an objective social reality that can be measured through observable, quantifiable indicators and analysed using statistical methods to identify law-like relationships (Saunders et al., 2019). This paradigm is consistent with the study's reliance on established acceptance theories that posit measurable causal relationships between perceptions and behaviour (Davis, 1989). The positivist stance underpins the use of a structured instrument, a large sample, and inferential statistics, and it positions the researcher as an objective observer external to the phenomenon under study.")

heading2(doc, "3.4 Inductive versus Deductive Reasoning")
body_para(doc,
    "The study employs deductive reasoning. It begins with established theory — the Technology Acceptance Model and its extensions — from which specific, testable hypotheses are derived and then evaluated against empirical data (Saunders et al., 2019). This contrasts with an inductive logic, which would build theory from observation. The deductive approach is appropriate because the study's purpose is to test the applicability of well-developed theory to a new population and context rather than to generate new theory, and because the constructs and their hypothesised relationships are already clearly specified in the literature (Venkatesh et al., 2003).")

heading2(doc, "3.5 Time Horizon")
body_para(doc,
    "The study adopts a cross-sectional time horizon, collecting data at a single point in time during the 2026 academic year (Saunders et al., 2019). A cross-sectional design is appropriate because the study aims to describe and analyse relationships among variables as they exist at a given moment, rather than to track change over time. This horizon is also practical within the constraints of an undergraduate dissertation, while still permitting robust correlational and regression analysis. The limitation that cross-sectional data cannot establish causal direction is acknowledged and addressed in the discussion of limitations.")

heading2(doc, "3.6 Research Strategy")
body_para(doc,
    "The research strategy is a survey, operationalised through a self-administered structured questionnaire (Saunders et al., 2019). The survey strategy is well suited to collecting standardised data from a large number of respondents efficiently and to producing data amenable to statistical analysis. It aligns with the positivist paradigm and deductive logic, and it permits the measurement of all study constructs using validated Likert-scale items adapted from prior adoption research (Davis, 1989; Strzelecki, 2023). The survey will be administered both in person and through an online survey tool to maximise reach and response rates.")

heading2(doc, "3.7 Sampling Frame and Sample Size")
body_para(doc,
    "The target population comprises undergraduate accounting and finance students at ZCAS University and the University of Zambia. The sampling frame consists of enrolment lists for accounting and finance programmes at the two institutions. The study employs stratified random sampling, stratifying by institution and year of study to ensure representation across both universities and all year groups, with random selection within strata to reduce selection bias (Saunders et al., 2019). Sample size is determined using the Krejcie and Morgan (1970) table, which provides recommended sample sizes for given population sizes at a 95% confidence level and 5% margin of error. For a target population in the low thousands, the table recommends a sample in the region of 300; recognising the practical constraints of an undergraduate study, a working sample of approximately 100 to 150 respondents is adopted, which is adequate for multiple regression with six predictors while acknowledging the trade-off in precision. The achieved sample and its adequacy will be reported transparently.")

heading2(doc, "3.8 Data Collection")
body_para(doc,
    "Primary data will be collected using a structured questionnaire comprising eight sections. Section A captures respondent profile information, while Sections B to H measure awareness, perceived usefulness, perceived ease of use, digital literacy/self-efficacy, institutional support, social/peer influence, and adoption behaviour, each using a five-point Likert scale ranging from 1 (strongly disagree) to 5 (strongly agree). The items are adapted from validated instruments in the technology-acceptance literature to ensure measurement quality (Davis, 1989; Venkatesh et al., 2003; Strzelecki, 2023). A pilot test will be conducted with a small group of students to refine wording and confirm clarity before full distribution. Questionnaires will be administered in person on campus and via an online survey platform, with informed consent obtained at the point of participation.")

heading2(doc, "3.9 Data Processing and Analysis")
body_para(doc,
    "Completed questionnaires will be screened for completeness, coded, and entered into statistical software for analysis. Data analysis will proceed in two stages (Saunders et al., 2019). Descriptive statistics, including frequencies, percentages, means, and standard deviations, will summarise the respondent profile and the distribution of construct scores. Inferential analysis will use Pearson correlation to assess bivariate relationships and multiple regression to test the combined and relative influence of the six independent variables on adoption, thereby evaluating the study's hypotheses. Moderation by year of study, gender, and prior technology experience will be examined through subgroup comparison or interaction terms. Results will be presented in tables and interpreted in relation to the research questions and prior literature.")

heading2(doc, "3.10 Reliability")
body_para(doc,
    "Reliability refers to the consistency of measurement (Saunders et al., 2019). Internal-consistency reliability will be assessed using Cronbach's alpha for each multi-item construct, with a minimum acceptable threshold of 0.70. Items that reduce a scale's reliability will be considered for removal following item-total correlation analysis. The use of established, previously validated scales further supports reliability, as does the pilot test, which allows identification and correction of ambiguous items before full administration (Davis, 1989; Strzelecki, 2023).")

heading2(doc, "3.11 Validity")
body_para(doc,
    "Validity concerns whether the instrument measures what it purports to measure (Saunders et al., 2019). Content validity will be established through expert review, in which the supervisor and subject specialists evaluate the relevance and representativeness of the items against the constructs. Construct validity will be examined through exploratory factor analysis to confirm that items load onto their intended constructs and to assess convergent and discriminant validity. The adaptation of items from instruments with demonstrated validity in prior adoption research provides an additional foundation for validity (Venkatesh et al., 2003).")

heading2(doc, "3.12 Generalisability")
body_para(doc,
    "Generalisability, or external validity, concerns the extent to which findings can be extended beyond the study sample (Saunders et al., 2019). Because the study samples accounting and finance students at two Lusaka-based institutions using stratified random sampling, findings are most directly generalisable to comparable urban Zambian higher-education contexts. Caution is warranted in extending the findings to rural institutions or to disciplines outside accounting and finance, given documented differences in infrastructure and disciplinary culture (International Energy Agency, 2024; Mutelo, 2025). The study's transparent reporting of sampling and response rates allows readers to assess transferability.")

heading2(doc, "3.13 Ethical and Access Issues")
body_para(doc,
    "The study will observe established ethical principles (Saunders et al., 2019). Ethical approval will be sought from the relevant institutional authorities at ZCAS University and UNZA, and access to enrolment frames will be requested through appropriate channels. Participation will be voluntary, with informed consent obtained before data collection; respondents will be assured of anonymity and confidentiality, and data will be stored securely and used solely for academic purposes. No personally identifying information will be reported, and respondents will be free to withdraw at any point without penalty. These measures protect participants and uphold the integrity of the research.")

heading2(doc, "3.14 Chapter Summary")
body_para(doc,
    "This chapter set out a positivist, deductive, quantitative methodology operationalised through a cross-sectional survey of accounting and finance students at ZCAS and UNZA. It justified stratified random sampling, a sample of approximately 100 to 150 respondents informed by the Krejcie and Morgan (1970) table, and a structured Likert-scale questionnaire analysed through descriptive statistics and multiple regression. Reliability, validity, generalisability, and ethical safeguards were addressed. The next chapter presents the data-analysis plan and template tables to be populated once primary data are collected.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER FOUR
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "CHAPTER FOUR: DATA ANALYSIS PLAN AND PRESENTATION")

heading2(doc, "4.0 Introduction")
body_para(doc,
    "This chapter will present and interpret the data collected through the questionnaire described in Chapter Three. As primary data have not yet been collected, this chapter currently provides the analysis plan and template table shells. No data have been fabricated. Once collected, the data are expected — in line with the international and Zambian literature reviewed in Chapter Two — to support the alternative hypothesis (H1), namely that the identified factors positively influence adoption (Strzelecki, 2023; Mudenda et al., 2026).")

heading2(doc, "4.1 Analysis Plan")
body_para(doc, "Analysis will proceed in four stages:")
numbered_para(doc, "The respondent profile will be summarised using frequencies and percentages to describe the sample composition by gender, age, year of study, institution, programme, device ownership, and internet access.")
numbered_para(doc, "Descriptive statistics will characterise each construct using means, standard deviations, and the range of scores to identify the distribution of perceptions and reported behaviour.")
numbered_para(doc, "Reliability will be assessed using Cronbach's alpha (threshold: α > 0.70) and validity examined through exploratory factor analysis to confirm scale integrity before inferential analysis.")
numbered_para(doc, "Pearson correlation and multiple regression will test the hypothesised relationships, with moderation by year of study, gender, and prior technology experience examined through subgroup comparisons or interaction terms.")

heading2(doc, "4.2 Template Table Shells")
body_para(doc, "The following tables provide the structure for data to be entered once primary data are collected. Dashes indicate cells awaiting empirical data.")
add_blank(doc)

# Table 4.1
mixed_para(doc, [("Table 4.1: ", True, False), ("Respondent Demographic Frequencies", False, True)])
add_table(doc,
    headers=["Variable", "Category", "Frequency (n)", "Percentage (%)"],
    rows=[
        ["Gender", "Male", "—", "—"],
        ["", "Female", "—", "—"],
        ["", "Prefer not to say", "—", "—"],
        ["Age range", "Under 20", "—", "—"],
        ["", "20–24", "—", "—"],
        ["", "25–29", "—", "—"],
        ["", "30 and above", "—", "—"],
        ["Year of study", "Year 1", "—", "—"],
        ["", "Year 2", "—", "—"],
        ["", "Year 3", "—", "—"],
        ["", "Year 4", "—", "—"],
        ["Institution", "ZCAS University", "—", "—"],
        ["", "UNZA", "—", "—"],
        ["", "Other", "—", "—"],
        ["Device ownership", "Smartphone only", "—", "—"],
        ["", "Laptop only", "—", "—"],
        ["", "Both", "—", "—"],
        ["Regular internet access", "Yes", "—", "—"],
        ["", "No", "—", "—"],
    ],
    col_widths=[4.0, 4.5, 3.5, 3.5])

add_blank(doc)
body_para(doc, "Note. Frequencies and percentages to be entered following data collection.", italic=True, size=10)
add_blank(doc)

# Table 4.2
mixed_para(doc, [("Table 4.2: ", True, False), ("Descriptive Statistics of Constructs", False, True)])
add_table(doc,
    headers=["Construct", "No. of Items", "Mean", "SD", "Min", "Max"],
    rows=[
        ["Awareness of AI Tools", "5", "—", "—", "—", "—"],
        ["Perceived Usefulness", "5", "—", "—", "—", "—"],
        ["Perceived Ease of Use", "5", "—", "—", "—", "—"],
        ["Digital Literacy / Self-Efficacy", "5", "—", "—", "—", "—"],
        ["Institutional Support", "5", "—", "—", "—", "—"],
        ["Social / Peer Influence", "4", "—", "—", "—", "—"],
        ["Adoption Behaviour", "5", "—", "—", "—", "—"],
    ],
    col_widths=[4.8, 2.5, 2.0, 2.0, 2.0, 2.0])

add_blank(doc)
body_para(doc, "Note. Likert scale: 1 = Strongly Disagree to 5 = Strongly Agree. Values to be entered following data collection.", italic=True, size=10)
add_blank(doc)

# Table 4.3
mixed_para(doc, [("Table 4.3: ", True, False), ("Reliability Results (Cronbach's Alpha)", False, True)])
add_table(doc,
    headers=["Construct", "No. of Items", "Cronbach's α", "Decision (α ≥ 0.70)"],
    rows=[
        ["Awareness of AI Tools", "5", "—", "—"],
        ["Perceived Usefulness", "5", "—", "—"],
        ["Perceived Ease of Use", "5", "—", "—"],
        ["Digital Literacy / Self-Efficacy", "5", "—", "—"],
        ["Institutional Support", "5", "—", "—"],
        ["Social / Peer Influence", "4", "—", "—"],
        ["Adoption Behaviour", "5", "—", "—"],
    ],
    col_widths=[4.8, 2.5, 3.0, 4.0])

add_blank(doc)
body_para(doc, "Note. Alpha coefficients to be computed from collected data. Items with item-total correlations below 0.30 will be considered for removal.", italic=True, size=10)
add_blank(doc)

# Table 4.4
mixed_para(doc, [("Table 4.4: ", True, False), ("Pearson Correlation Matrix", False, True)])
constructs_short = ["Awareness", "Perc. Usefulness", "Perc. Ease", "Digital Lit.", "Inst. Support", "Peer Influence", "Adoption"]
corr_rows = []
for i, c in enumerate(constructs_short):
    row = [c] + ["—"]*7
    row[i+1] = "1.00"
    corr_rows.append(row)
add_table(doc,
    headers=["Construct", "1", "2", "3", "4", "5", "6", "7"],
    rows=[[f"{i+1}. {c}"] + ["—"]*7 for i, c in enumerate(constructs_short)],
    col_widths=[4.0, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4])

add_blank(doc)
body_para(doc, "Note. Pearson r coefficients to be entered following data collection. * p < .05, ** p < .01.", italic=True, size=10)
add_blank(doc)

# Table 4.5
mixed_para(doc, [("Table 4.5: ", True, False), ("Multiple Regression Results — Predictors of AI Adoption", False, True)])
add_table(doc,
    headers=["Predictor", "B", "SE", "β", "t", "p"],
    rows=[
        ["(Constant)", "—", "—", "—", "—", "—"],
        ["Perceived Usefulness", "—", "—", "—", "—", "—"],
        ["Perceived Ease of Use", "—", "—", "—", "—", "—"],
        ["Digital Literacy / Self-Efficacy", "—", "—", "—", "—", "—"],
        ["Institutional Support", "—", "—", "—", "—", "—"],
        ["Social / Peer Influence", "—", "—", "—", "—", "—"],
        ["Awareness of AI Tools", "—", "—", "—", "—", "—"],
        ["R²", "", "", "", "", "—"],
        ["Adjusted R²", "", "", "", "", "—"],
        ["F-statistic", "", "", "", "", "—"],
    ],
    col_widths=[5.0, 1.8, 1.8, 1.8, 1.8, 1.8])

add_blank(doc)
body_para(doc, "Note. B = unstandardised coefficient; SE = standard error; β = standardised coefficient. Values to be computed following data collection. p < .05 denotes statistical significance.", italic=True, size=10)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER FIVE
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "CHAPTER FIVE: CONCLUSIONS")

heading2(doc, "5.0 Introduction")
body_para(doc,
    "This chapter draws conclusions from the study, organised by research objective. As primary data have not yet been collected, the conclusions presented are predictions grounded in the literature reviewed in Chapter Two and lean toward supporting the study's hypothesis; they are to be revised once empirical findings are available. The chapter then discusses practical and managerial implications, acknowledges limitations, identifies directions for future research, and offers concluding remarks.")

heading2(doc, "5.1 Conclusion on Objective One: Perceived Usefulness and Perceived Ease of Use")
body_para(doc,
    "The literature predicts that perceived usefulness and perceived ease of use will emerge as significant positive determinants of AI adoption among accounting and finance students, with perceived usefulness likely the stronger of the two (Davis, 1989; Strzelecki, 2023). Given that Zambian students already report high adoption rates (Chaamwe, 2025; Mudenda et al., 2026), the study anticipates that students who perceive AI tools as enhancing performance and as effortless to use will report greater adoption. This conclusion will be confirmed or qualified by the regression results once primary data are collected and analysed.")

heading2(doc, "5.2 Conclusion on Objective Two: Digital Literacy and Self-Efficacy")
body_para(doc,
    "The study anticipates that digital literacy and self-efficacy will positively influence adoption, consistent with evidence that computer self-efficacy is a robust predictor of technology acceptance (Hatlevik, Throndsen, Loi, & Gudmundsdóttir, 2018). Students with greater confidence in their digital skills are expected to adopt AI tools more readily. Should self-efficacy prove influential in the regression results, it would indicate that AI-literacy training could meaningfully raise adoption among less confident students, providing a targeted and actionable lever for institutions.")

heading2(doc, "5.3 Conclusion on Objective Three: Institutional Support and Social/Peer Influence")
body_para(doc,
    "Drawing on UTAUT, the study expects that institutional support (facilitating conditions) and social or peer influence will positively affect adoption, though their relative weight may vary (Venkatesh, Morris, Davis, & Davis, 2003; Strzelecki, 2023). In the Zambian context, where electricity access and connectivity remain constrained, institutional support may prove especially consequential (International Energy Agency, 2024; Mutelo, 2025). Peer influence is expected to matter given the social diffusion of AI tools already documented among Zambian students.")

heading2(doc, "5.4 Conclusion on Objective Four: Awareness and Moderating Variables")
body_para(doc,
    "The study anticipates that awareness of AI tools will positively predict adoption, consistent with African evidence identifying awareness as a precursor to use (Mudenda et al., 2026). The moderating variables — year of study, gender, and prior technology experience — are expected to condition relationships, with prior experience likely strengthening the effect of ease of use and gender or year of study potentially moderating social and institutional effects (Venkatesh et al., 2003). These moderation effects will be tested empirically and will determine whether targeted interventions should be differentiated by student profile.")

heading2(doc, "5.5 Practical and Managerial Implications")
body_para(doc,
    "If the anticipated findings are confirmed, several implications follow for institutions, lecturers, and policymakers. Institutions such as ZCAS and UNZA should invest in reliable connectivity and device access to address facilitating conditions, recognising that infrastructure is a binding constraint in the Zambian higher-education context (International Energy Agency, 2024). AI-literacy training should be integrated into accounting and finance curricula to build self-efficacy among students who may lack confidence in using new technologies. Clear institutional policies should guide responsible AI use while safeguarding academic integrity in assessment-intensive programmes (ACCA, 2025; Yusuf, Pervin, & Román-González, 2024). Lecturers should model constructive AI use and redesign assessments to reduce opportunities for misconduct without stifling the legitimate productivity benefits of AI tools. Policymakers can align institutional initiatives with the national digital-skills agenda articulated in the National Digital Transformation Strategy 2023–2027 (Government of the Republic of Zambia, 2023).")

heading2(doc, "5.6 Limitations and Future Research")
body_para(doc,
    "The study is limited by several factors that future research should address. The cross-sectional design precludes causal inference; relationships identified through regression establish association rather than direction. Concentration on two Lusaka-based institutions limits generalisability to rural and peri-urban contexts where infrastructure constraints may produce different adoption patterns. Reliance on self-reported data may introduce social desirability bias, particularly regarding academic integrity. The modest sample size relative to the total population reduces statistical power for detecting smaller effects in the regression and moderation analyses.")
body_para(doc,
    "Future research could employ longitudinal designs to capture change in adoption over time as AI tools evolve. Expansion to multiple institutions, including rural universities, would broaden the generalisability of findings. Incorporating qualitative methods would illuminate the reasoning and motivations behind adoption decisions that structured surveys cannot capture. A dedicated study examining the relationship between academic-integrity concerns and adoption decisions among accounting and finance students would address a gap identified in the literature.")

heading2(doc, "5.7 Conclusions")
body_para(doc,
    "The study is positioned to establish that the adoption of AI-powered learning technologies among accounting and finance students in Zambia is shaped by a combination of perceptual, individual, institutional, and social factors, consistent with the Technology Acceptance Model and its extensions. By isolating accounting and finance students within the Zambian context, the study contributes disciplinary and contextual specificity that the existing literature lacks, offering evidence to guide institutional and national efforts to integrate AI responsibly into higher education. The study's findings, once primary data are collected and analysed, will either confirm the alternative hypothesis — that all six identified factors positively influence adoption — or reveal the relative weighting of factors and the conditions under which adoption is constrained.")

heading2(doc, "5.8 Chapter Summary")
body_para(doc,
    "This chapter presented literature-grounded predictive conclusions for each objective, discussed practical implications for institutions, lecturers, and policymakers, acknowledged methodological limitations, and identified avenues for future research. These conclusions are provisional and will be revised once primary data are collected and analysed. The complete questionnaire instrument and the APA 7th edition reference list follow.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "REFERENCES")
add_blank(doc)

refs = [
    "Association of Chartered Certified Accountants. (2025). AI monitor: How artificial intelligence is reshaping the accountancy profession. ACCA Global.",
    "Aga, M. K. (2025). Accounting students' technology readiness, perceptions, and digital competence toward artificial intelligence adoption in accounting curricula. Journal of Accounting Education, 70, Article 100968. https://doi.org/10.1016/j.jaccedu.2025.100968",
    "Al-Adwan, A. S., Li, N., Al-Adwan, A., Abbasi, G. A., Albelbisi, N. A., & Habibi, A. (2023). Extending the technology acceptance model (TAM) to predict university students' intentions to use metaverse-based learning platforms. Education and Information Technologies, 28(11), 15381–15413. https://doi.org/10.1007/s10639-023-11816-3",
    "Bin-Nashwan, S. A., Sadallah, M., & Bouteraa, M. (2023). Use of ChatGPT in academia: Academic integrity hangs in the balance. Technology in Society, 75, Article 102370. https://doi.org/10.1016/j.techsoc.2023.102370",
    "Bonsu, E. M., & Baffour-Koduah, D. (2023). From the consumers' side: Determining students' perception and intention to use ChatGPT in Ghanaian higher education. Journal of Education, Society and Multiculturalism, 4(1), 1–29. https://doi.org/10.2478/jesm-2023-0001",
    "Chaamwe, N. (2025). Investigating the factors influencing students' adoption of generative AIs in universities: A case of the Copperbelt University. Zambia ICT Journal, 8(1), 47–53. https://doi.org/10.33260/zictjournal.v8i1.340",
    "Compeau, D. R., & Higgins, C. A. (1995). Computer self-efficacy: Development of a measure and initial test. MIS Quarterly, 19(2), 189–211. https://doi.org/10.2307/249688",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340. https://doi.org/10.2307/249008",
    'Dwivedi, Y. K., Kshetri, N., Hughes, L., Slade, E. L., Jeyaraj, A., Kar, A. K., & Wright, R. (2023). Opinion paper: “So what if ChatGPT wrote it?” Multidisciplinary perspectives on opportunities, challenges and implications of generative conversational AI for research, practice and policy. International Journal of Information Management, 71, Article 102642. https://doi.org/10.1016/j.ijinfomgt.2023.102642',
    "Government of the Republic of Zambia. (2023). National information and communication technology (ICT) policy 2023. Ministry of Technology and Science.",
    "Hatlevik, O. E., Throndsen, I., Loi, M., & Gudmundsdóttir, G. B. (2018). Students' ICT self-efficacy and computer and information literacy: Determinants and relationships. Computers & Education, 118, 107–119. https://doi.org/10.1016/j.compedu.2017.11.011",
    "International Energy Agency. (2024). SDG7: Data and projections — Access to electricity. IEA. https://www.iea.org/reports/sdg7-data-and-projections",
    "Kasneci, E., Sessler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, Article 102274. https://doi.org/10.1016/j.lindif.2023.102274",
    "Krejcie, R. V., & Morgan, D. W. (1970). Determining sample size for research activities. Educational and Psychological Measurement, 30(3), 607–610. https://doi.org/10.1177/001316447003000308",
    "Ma, W., Adesope, O. O., Nesbit, J. C., & Liu, Q. (2014). Intelligent tutoring systems and learning outcomes: A meta-analysis. Journal of Educational Psychology, 106(4), 901–918. https://doi.org/10.1037/a0037123",
    "Mudenda, S., Mukosha, M., Mfune, R. L., Kathewera, B., Mutanekelwa, I., Mwanza, B., Mufwambi, W., Hampango, M., Kamvuma, K., Mwaba, M., Muyenga, T., Chileshe, C., Zulu, M., Tembo, R., Mwaba, F., Kafwimbi, S., Lubanga, A. F., Simweene, C. C., Mohamed, S., … Godman, B. (2026). Integrating generative artificial intelligence in African higher education: University students' awareness, attitudes, and use of ChatGPT in Zambia. Frontiers in Education, 11, Article 1814033. https://doi.org/10.3389/feduc.2026.1814033",
    "Mutelo, I. (2025). Understanding the generative artificial intelligence revolution in Zambian higher education research: Adoption, challenges, and strategies for responsible integration. International Journal of Research and Innovation in Social Science, 9(3S), 5731–5737. https://doi.org/10.47772/IJRISS.2025.903SEDU0416",
    "Okello, F. (2023). Bridging Kenya's digital divide: Context, barriers and strategies. Centre for International Governance Innovation.",
    "Polyportis, A. (2024). A longitudinal study on artificial intelligence adoption: Understanding the drivers of ChatGPT usage behavior change in higher education. Frontiers in Artificial Intelligence, 6, Article 1324398. https://doi.org/10.3389/frai.2023.1324398",
    "Raman, R., Mandal, S., Das, P., Kaur, T., Sanjanasri, J. P., & Nedungadi, P. (2024). Exploring university students' adoption of ChatGPT using the diffusion of innovation theory and sentiment analysis with gender dimension. Human Behavior and Emerging Technologies, 2024, Article 3085910. https://doi.org/10.1155/2024/3085910",
    "Rogers, E. M. (2003). Diffusion of innovations (5th ed.). Free Press.",
    "Saunders, M. N. K., Lewis, P., & Thornhill, A. (2019). Research methods for business students (8th ed.). Pearson Education.",
    "Shahzad, M. F., Xu, S., & Asif, M. (2025). Factors affecting generative artificial intelligence, such as ChatGPT, use in higher education: An application of technology acceptance model. British Educational Research Journal. Advance online publication. https://doi.org/10.1002/berj.4084",
    "Steenbergen-Hu, S., & Cooper, H. (2014). A meta-analysis of the effectiveness of intelligent tutoring systems on college students' academic learning. Journal of Educational Psychology, 106(2), 331–347. https://doi.org/10.1037/a0034752",
    "Strzelecki, A. (2023). To use or not to use ChatGPT in higher education? A study of students' acceptance and use of technology. Interactive Learning Environments, 32(9), 5142–5155. https://doi.org/10.1080/10494820.2023.2209881",
    "Strzelecki, A., & ElArabawy, S. (2024). Investigation of the moderation effect of gender and study level on the acceptance and use of generative AI by higher education students: Comparative evidence from Poland and Egypt. British Journal of Educational Technology, 55(3), 1209–1230. https://doi.org/10.1111/bjet.13425",
    "Tiwari, C. K., Bhat, M. A., Khan, S. T., Subramaniam, R., & Khan, M. A. I. (2023). What drives students toward ChatGPT? An investigation of the factors influencing adoption and usage of ChatGPT. Interactive Technology and Smart Education, 21(3), 333–355. https://doi.org/10.1108/ITSE-04-2023-0061",
    "Tran, Q. T. T. (2025). Impact of AI adoption on the digital workplace and accounting education for university students. Industry and Higher Education. Advance online publication. https://doi.org/10.1177/09504222251376760",
    "Venkatesh, V., & Davis, F. D. (2000). A theoretical extension of the technology acceptance model: Four longitudinal field studies. Management Science, 46(2), 186–204. https://doi.org/10.1287/mnsc.46.2.186.11926",
    "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425–478. https://doi.org/10.2307/30036540",
    "Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer acceptance and use of information technology: Extending the unified theory of acceptance and use of technology. MIS Quarterly, 36(1), 157–178. https://doi.org/10.2307/41410412",
    "Yusuf, A., Pervin, N., & Román-González, M. (2024). Generative AI and the future of higher education: A threat to academic integrity or reformation? Evidence from multicultural perspectives. International Journal of Educational Technology in Higher Education, 21, Article 21. https://doi.org/10.1186/s41239-024-00453-6",
]

for ref in refs:
    # Hanging indent for APA references
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    set_font(run)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX A: QUESTIONNAIRE
# ═══════════════════════════════════════════════════════════════════════════

heading1(doc, "APPENDIX A: RESEARCH QUESTIONNAIRE")
add_blank(doc)

mixed_para(doc, [("Title: ", True, False), ("Factors Influencing the Adoption of AI-Powered Learning Technologies Among Accounting and Finance Students in Zambia", False, False)])
add_blank(doc)

body_para(doc,
    "Thank you for agreeing to participate in this study. This questionnaire forms part of an undergraduate dissertation at ZCAS University. Your participation is voluntary and your responses are anonymous and confidential. The questionnaire takes approximately 10–12 minutes to complete.")
add_blank(doc)
body_para(doc,
    '"AI-powered learning technologies" refers to tools such as ChatGPT and other generative AI applications, AI features in learning platforms (e.g., Coursera), and intelligent tutoring or adaptive learning systems used for studying.')
add_blank(doc)
mixed_para(doc, [
    ("For Sections B–H, please indicate your level of agreement with each statement by circling or selecting one number, where ", False, False),
    ("1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree, 5 = Strongly Agree.", True, False),
    (" Please answer honestly; there are no right or wrong answers.", False, False),
])
add_blank(doc)

# Section A
heading2(doc, "Section A: Respondent Profile (please tick one)")

q_a = [
    ("A1", "Gender:", "☐ Male     ☐ Female     ☐ Prefer not to say"),
    ("A2", "Age range:", "☐ Under 20     ☐ 20–24     ☐ 25–29     ☐ 30 and above"),
    ("A3", "Year of study:", "☐ Year 1     ☐ Year 2     ☐ Year 3     ☐ Year 4"),
    ("A4", "Institution:", "☐ ZCAS University     ☐ University of Zambia (UNZA)     ☐ Other (specify): _______________"),
    ("A5", "Programme:", "☐ Accounting     ☐ Finance     ☐ Financial Services     ☐ Other (specify): _______________"),
    ("A6", "Do you own a smartphone?", "☐ Yes     ☐ No"),
    ("A7", "Do you own a laptop?", "☐ Yes     ☐ No"),
    ("A8", "Do you have regular access to the internet?", "☐ Yes     ☐ No"),
]

for code, question, options in q_a:
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=2)
    r1 = p.add_run(f"{code}. {question} ")
    set_font(r1, bold=True)
    r2 = p.add_run(options)
    set_font(r2)

add_blank(doc)

def add_section_likert(doc, section_letter, title, items):
    heading2(doc, f"Section {section_letter}: {title} (1 = Strongly Disagree … 5 = Strongly Agree)")
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        set_spacing(p, before=4, after=2)
        r1 = p.add_run(f"{section_letter}{i}. ")
        set_font(r1, bold=True)
        r2 = p.add_run(item)
        set_font(r2)
        # scale line
        p2 = doc.add_paragraph()
        set_spacing(p2, before=0, after=6)
        p2.paragraph_format.left_indent = Cm(0.5)
        r3 = p2.add_run("1 ☐     2 ☐     3 ☐     4 ☐     5 ☐")
        set_font(r3, size=10)
    add_blank(doc)

add_section_likert(doc, "B", "Awareness of AI-Powered Learning Technologies", [
    "I am aware that AI-powered learning technologies exist and can be used for studying.",
    "I know of specific AI tools (such as ChatGPT) that can assist with my coursework.",
    "I am aware of how AI tools can be applied to accounting and finance learning.",
    "I have seen or heard AI learning tools discussed at my institution.",
    "I am well informed about the range of AI learning tools currently available.",
])

add_section_likert(doc, "C", "Perceived Usefulness", [
    "Using AI-powered learning technologies improves my academic performance.",
    "AI tools help me understand accounting and finance concepts more quickly.",
    "AI tools increase my productivity when completing assignments.",
    "Using AI tools makes my studying more effective.",
    "Overall, I find AI-powered learning technologies useful for my studies.",
])

add_section_likert(doc, "D", "Perceived Ease of Use", [
    "Learning to use AI-powered learning technologies is easy for me.",
    "I find AI tools clear and understandable to interact with.",
    "It is easy to become skilful at using AI learning tools.",
    "I can use AI tools without needing technical assistance.",
    "Overall, I find AI-powered learning technologies easy to use.",
])

add_section_likert(doc, "E", "Digital Literacy / Self-Efficacy", [
    "I am confident in my ability to use digital technologies for learning.",
    "I can solve most technical problems I encounter when using digital tools.",
    "I have the digital skills needed to use AI learning tools effectively.",
    "I feel comfortable trying new technologies on my own.",
    "I am confident I could use a new AI learning tool even without prior training.",
])

add_section_likert(doc, "F", "Institutional Support", [
    "My institution provides adequate internet access for using AI learning tools.",
    "My lecturers encourage the responsible use of AI learning tools.",
    "My institution provides guidance or training on using AI tools for learning.",
    "The infrastructure at my institution (devices, electricity, connectivity) supports AI tool use.",
    "My institution has clear policies on the use of AI in learning and assessment.",
])

add_section_likert(doc, "G", "Social / Peer Influence", [
    "My fellow students use AI-powered learning technologies for their studies.",
    "People whose opinions I value think I should use AI learning tools.",
    "My classmates encourage me to use AI tools for coursework.",
    "Using AI learning tools is common practice among students in my programme.",
])

add_section_likert(doc, "H", "Actual / Intended Adoption Behaviour (Dependent Variable)", [
    "I currently use AI-powered learning technologies for my studies.",
    "I intend to use AI learning tools regularly in the future.",
    "I use AI tools to support my accounting and finance coursework.",
    "I would recommend AI learning tools to other students.",
    "I plan to continue using AI-powered learning technologies throughout my studies.",
])

add_blank(doc)
body_para(doc, "Thank you for your participation. Your responses are greatly appreciated and will contribute to improving AI integration in Zambian higher education.",
          align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved -> {OUT}")
